# --- START OF FILE footnotes_service.py ---

# services/footnotes_service.py

import docx
from docx.enum.text import WD_COLOR_INDEX
from docx.shared import RGBColor, Pt
import re
from collections import Counter, defaultdict
import logging
import time
import copy
from .common_docx import _safe_copy_paragraph_format, _apply_run_style, tokenize_paragraph_universal

# --- Настройка логгера ---
logger_fn = logging.getLogger(__name__)

# --- Сервисы ---
try:
    # Импортируем функции из pymorphy_service, необходимые для этого сервиса
    from services.pymorphy_service import (
        load_pymorphy, load_nltk_lemmatizer, _get_lemma, _get_stem,
        WORD_TOKENIZE_PATTERN

    )
    logger_fn.debug("Функции из pymorphy_service успешно импортированы.")
except ImportError:
    logger_fn.error("Не удалось импортировать pymorphy_service. Используются заглушки.", exc_info=True)
    load_pymorphy=lambda:None
    load_nltk_lemmatizer=lambda:None
    _get_lemma=lambda w,uc=True: w.lower() if w else None
    _get_stem=lambda w,uc=True: w.lower() if w else None
    log_cache_stats=lambda:None
    reset_cache_stats=lambda:None
    WORD_TOKENIZE_PATTERN=re.compile(r'\b\w+\b', re.UNICODE)


# --- КОНСТАНТЫ ---
HIGHLIGHT_COLOR = WD_COLOR_INDEX.BRIGHT_GREEN # Цвет выделения


# --- Поиск совпадений (переименовано, логика без изменений) ---
def _find_matches_for_highlighting(
    tokens,
    search_lemmas_map,           # Карта для слов по леммам {lemma: info}
    search_stems_map,            # Карта для слов по стеммам {stem: info}
    search_phrase_lemmas_map,    # Карта для фраз по леммам {tuple(sorted_lemmas): info}
    search_phrase_stems_map      # Карта для фраз по стеммам {tuple(sorted_stems): info}
):
    """
    Находит совпадения по ЛЕММАМ или СТЕММАМ, БЕЗ УЧЕТА ПОРЯДКА СЛОВ во фразах.
    Приоритет: фразы > слова, длинные > короткие, леммы > стеммы.
    Разрешает перекрытия в пользу лучших совпадений.
    """
    potential_matches = []
    word_tokens_with_indices = [
        (idx, t) for idx, t in enumerate(tokens)
        if t['type'] == 'word' and (t['lemma'] is not None or t['stem'] is not None)
    ]
    if not word_tokens_with_indices: return []
    num_word_tokens = len(word_tokens_with_indices)

    # 1. Поиск ФРАЗ (без учета порядка)
    phrase_maps_available = search_phrase_lemmas_map or search_phrase_stems_map
    if phrase_maps_available:
        all_phrase_lengths = set()
        if search_phrase_lemmas_map: all_phrase_lengths.update(len(p_lemmas) for p_lemmas in search_phrase_lemmas_map.keys())
        if search_phrase_stems_map: all_phrase_lengths.update(len(p_stems) for p_stems in search_phrase_stems_map.keys())
        sorted_phrase_lengths = sorted(list(all_phrase_lengths), reverse=True)

        if sorted_phrase_lengths:
            min_phrase_len = sorted_phrase_lengths[-1]
            for i in range(num_word_tokens - min_phrase_len + 1):
                for phrase_len in sorted_phrase_lengths:
                    if i + phrase_len > num_word_tokens: continue
                    current_word_tokens_info = word_tokens_with_indices[i : i + phrase_len]
                    first_token_info = current_word_tokens_info[0]
                    last_token_info = current_word_tokens_info[-1]

                    current_lemmas_list = [info[1]['lemma'] for info in current_word_tokens_info if info[1]['lemma']]
                    current_stems_list = [info[1]['stem'] for info in current_word_tokens_info if info[1]['stem']]
                    sorted_lemmas_from_text = tuple(sorted(current_lemmas_list))
                    sorted_stems_from_text = tuple(sorted(current_stems_list))

                    found_match_info = None; match_method = None
                    if search_phrase_lemmas_map and sorted_lemmas_from_text in search_phrase_lemmas_map:
                        found_match_info = search_phrase_lemmas_map[sorted_lemmas_from_text]; match_method = 'lemma'
                    elif search_phrase_stems_map and sorted_stems_from_text in search_phrase_stems_map:
                        found_match_info = search_phrase_stems_map[sorted_stems_from_text]; match_method = 'stem'

                    if found_match_info:
                        potential_matches.append({
                            'start': first_token_info[1]['start'],
                            'end': last_token_info[1]['end'],
                            'match_type': 'phrase',
                            'lemma_key': found_match_info['lemma_key'], # Ключ лемм из поиска
                            'original_form': found_match_info['original'],
                            'length': phrase_len,
                            'match_method': match_method,
                            'token_indices': {info[0] for info in current_word_tokens_info}
                        })

    # 2. Поиск ОДИНОЧНЫХ СЛОВ
    word_maps_available = search_lemmas_map or search_stems_map
    if word_maps_available:
        for word_token_index, word_token_data in word_tokens_with_indices:
            token_lemma = word_token_data['lemma']
            token_stem = word_token_data['stem']
            found_match_info = None; match_method = None
            if token_lemma and search_lemmas_map and token_lemma in search_lemmas_map:
                found_match_info = search_lemmas_map[token_lemma]; match_method = 'lemma'
            elif token_stem and search_stems_map and token_stem in search_stems_map:
                found_match_info = search_stems_map[token_stem]; match_method = 'stem'

            if found_match_info:
                potential_matches.append({
                    'start': word_token_data['start'],
                    'end': word_token_data['end'],
                    'match_type': 'word',
                    'lemma_key': found_match_info['lemma'], # Лемма из поиска
                    'original_form': found_match_info['original'],
                    'length': 1,
                    'match_method': match_method,
                    'token_indices': {word_token_index}
                })

    # 3. Фильтрация и Приоритезация совпадений
    final_matches = []
    covered_token_indices = set()

    potential_matches.sort(key=lambda x: (
        0 if x['match_type'] == 'phrase' else 1,
        -x['length'],
        0 if x['match_method'] == 'lemma' else 1,
        x['start']
    ))

    for match in potential_matches:
        if not match['token_indices'].intersection(covered_token_indices):
            final_matches.append(match)
            covered_token_indices.update(match['token_indices'])

    # 4. Финальная сортировка по начальной позиции
    final_matches.sort(key=lambda x: x['start'])
    return final_matches


# --- Воссоздание параграфа ТОЛЬКО с подсветкой (переименовано и упрощено) ---
def _reconstruct_paragraph_with_highlighting(
    new_paragraph, source_paragraph, tokens, matches,
    word_stats, phrase_stats # Параметры сносок убраны
):
    """Воссоздает параграф, добавляя ТОЛЬКО подсветку для найденных 'word' и 'phrase'."""
    # Копирование стиля/формата параграфа
    if source_paragraph.style and hasattr(source_paragraph.style, 'name'):
        try:
            if source_paragraph.style.name in new_paragraph.part.document.styles: new_paragraph.style = new_paragraph.part.document.styles[source_paragraph.style.name]
        except Exception as e: logger_fn.warning(f"Ошибка копирования стиля параграфа: {e}")
    if hasattr(source_paragraph, 'paragraph_format'): _safe_copy_paragraph_format(source_paragraph.paragraph_format, new_paragraph.paragraph_format)
    if not tokens: return

    current_pos = 0; paragraph_len = tokens[-1]['end'] if tokens else 0; source_runs = list(source_paragraph.runs); run_cache = {}

    # Функция поиска исходного run'а (без изменений)
    def find_run_at_or_after(char_pos):
        if char_pos in run_cache: return run_cache[char_pos]
        current_run_start = 0; best_run = None
        for run_index, run in enumerate(source_runs):
            run_len = len(run.text); run_end = current_run_start + run_len
            if current_run_start <= char_pos < run_end: best_run = run; break
            if char_pos < run_end:
                best_run = source_runs[run_index - 1] if run_index > 0 else run; break
            current_run_start = run_end
        else: best_run = source_runs[-1] if source_runs else None
        run_cache[char_pos] = best_run; return best_run

    # Итерация по совпадениям
    for match_idx, match in enumerate(matches):
        match_start = match['start']; match_end = match['end']; match_type = match['match_type']
        original_search_term = match['original_form'] # Оригинал из поиска
        lemma_key = match['lemma_key'] # Лемма/ключ лемм исходного термина
        match_method = match.get('match_method', 'unknown')

        # Добавляем текст ПЕРЕД совпадением
        if match_start > current_pos:
            segment_text = "".join(t['text'] for t in tokens if current_pos <= t['start'] < match_start)
            if segment_text:
                source_run_style = find_run_at_or_after(current_pos)
                new_run = new_paragraph.add_run(segment_text)
                if source_run_style: _apply_run_style(source_run_style, new_run, copy_highlight=False)

        # Добавляем текст СОВПАДЕНИЯ с подсветкой
        if match_end > match_start:
            segment_text = "".join(t['text'] for t in tokens if match_start <= t['start'] < match_end)
            if segment_text:
                source_run_style = find_run_at_or_after(match_start)
                # 1. Добавляем текст, копируем исходный стиль
                text_run = new_paragraph.add_run(segment_text)
                if source_run_style: _apply_run_style(source_run_style, text_run, copy_highlight=True) # Копируем и подсветку оригинала, если была
                # 2. Применяем НАШУ подсветку
                try: text_run.font.highlight_color = HIGHLIGHT_COLOR
                except Exception as e: logger_fn.error(f"Ошибка применения подсветки для '{segment_text}': {e}")

                # 3. Логика добавления сноски УДАЛЕНА

                # 4. Обновляем СТАТИСТИКУ (без номера сноски)
                found_text_in_doc = segment_text.strip() # Текст, как он найден в документе
                stats_entry = None; stats_key = None
                try:
                    if match_type == 'phrase':
                         stats_key = original_search_term # Ключ статистики - оригинальная фраза
                         stats_entry = phrase_stats[stats_key]
                    elif match_type == 'word':
                         stats_key = lemma_key # Ключ статистики - лемма исходного слова
                         if isinstance(stats_key, str): stats_entry = word_stats[stats_key]
                         else: logger_fn.warning(f"Некорректный ключ леммы: {stats_key}")

                    if stats_entry is not None:
                         stats_entry['count'] += 1
                         # Добавляем найденную форму в статистику
                         stats_entry['forms'][found_text_in_doc.lower() if match_type=='word' else found_text_in_doc] += 1
                         # Записываем оригинал, если его еще нет
                         if stats_entry.get('original') is None: stats_entry['original'] = original_search_term
                except Exception as e: logger_fn.error(f"Ошибка обновления статистики для '{stats_key}': {e}")

        # Обновляем текущую позицию
        current_pos = match_end

    # Добавляем текст ПОСЛЕ последнего совпадения
    if current_pos < paragraph_len:
        segment_text = "".join(t['text'] for t in tokens if t['start'] >= current_pos)
        if segment_text:
            source_run_style = find_run_at_or_after(current_pos)
            new_run = new_paragraph.add_run(segment_text)
            if source_run_style: _apply_run_style(source_run_style, new_run, copy_highlight=False)

    run_cache.clear() # Очищаем кэш для параграфа

# --- Основная функция анализа и ТОЛЬКО подсветки (переименована и упрощена) ---
def analyze_and_highlight_terms_docx(
    source_path,
    search_lemmas_map,
    search_stems_map,
    search_phrase_lemmas_map, # Ключ - отсортированный кортеж лемм
    search_phrase_stems_map,  # Ключ - отсортированный кортеж стеммов
    output_path
):
    """
    Анализирует документ, добавляет ТОЛЬКО подсветку для терминов, найденных
    по ЛЕММАМ или СТЕММАМ (фразы - без учета порядка слов).
    Собирает статистику найденных терминов.
    """
    start_process_time = time.time()
    logger_fn.info(f"Начало анализа (леммы+стеммы, фразы БЕЗ учета порядка) и ПОДСВЕТКИ: '{source_path}' -> '{output_path}'")
    logger_fn.info(f"Искомых лемм слов: {len(search_lemmas_map)}, стеммов слов: {len(search_stems_map)}")
    logger_fn.info(f"Искомых структур лемм фраз: {len(search_phrase_lemmas_map)}, стеммов фраз: {len(search_phrase_stems_map)}")

    try:
        # Инициализация и сброс кешей
        load_pymorphy(); load_nltk_lemmatizer()
        reset_cache_stats(); 

        # Структуры для статистики (без сносок)
        # found_original_term_to_footnote_number = {} # Убрано
        # footnote_counter = [1] # Убрано
        word_stats = defaultdict(lambda: {'count': 0, 'forms': Counter(), 'original': None}) # Ключ - лемма
        phrase_stats = defaultdict(lambda: {'count': 0, 'forms': Counter(), 'original': None}) # Ключ - оригинал фразы

        # --- Открытие документов ---
        logger_fn.info(f"Открытие исходного документа: {source_path}")
        try: doc = docx.Document(source_path)
        except Exception as e_open: logger_fn.error(f"Не удалось открыть '{source_path}': {e_open}"); return None
        result_doc = docx.Document() # Создаем новый документ для результата

        # --- Копирование стилей (без изменений) ---
        logger_fn.info("Копирование стилей..."); source_styles = doc.styles; target_styles = result_doc.styles; copied_count = 0; skipped_count = 0
        try:
            target_style_ids = {getattr(s, 'style_id', None) for s in target_styles}
            for style in source_styles:
                style_id = getattr(style, 'style_id', None)
                if hasattr(style, 'element') and style_id and style_id not in target_style_ids:
                    try: target_styles.element.append(copy.deepcopy(style.element)); target_style_ids.add(style_id); copied_count += 1
                    except Exception as e_copy_style: logger_fn.warning(f"Не удалось скопировать стиль '{getattr(style,'name','N/A')}': {e_copy_style}")
                elif style_id in target_style_ids: skipped_count +=1
        except Exception as e: logger_fn.error(f"Ошибка копирования стилей: {e}")
        logger_fn.info(f"Стили скопированы: {copied_count}, пропущено: {skipped_count}.")

        # --- Копирование свойств секции (без изменений) ---
        if doc.sections:
            try:
                source_sec = doc.sections[0]; target_sec = result_doc.sections[0]
                attrs_to_copy = ['orientation', 'page_height', 'page_width', 'left_margin', 'right_margin', 'top_margin', 'bottom_margin', 'header_distance', 'footer_distance', 'gutter']
                logger_fn.info("Копирование свойств секции...")
                for attr in attrs_to_copy:
                    try: val = getattr(source_sec, attr, None); setattr(target_sec, attr, val)
                    except Exception as e_attr: logger_fn.warning(f"Не удалось скопировать свойство секции '{attr}': {e_attr}")
            except Exception as e_sec: logger_fn.warning(f"Общая ошибка при копировании свойств секции: {e_sec}")

        # --- Обработка основного контента (параграфы и таблицы по порядку) ---
        body_elements = []
        try: # Безопасная итерация по телу документа
             for el in doc.element.body:
                 if el.tag.endswith('p'): body_elements.append({'type': 'paragraph', 'element': el})
                 elif el.tag.endswith('tbl'): body_elements.append({'type': 'table', 'element': el})
        except Exception as e_iter: logger_fn.error(f"Ошибка итерации по body: {e_iter}")

        logger_fn.info(f"Начинается обработка {len(body_elements)} элементов документа...")
        para_idx = 0; table_idx = 0 # Счетчики для логов
        # Итерация по сохраненным элементам
        for item in body_elements:
             el_type = item['type']; el = item['element']

             # --- Обработка Параграфа ---
             if el_type == 'paragraph':
                 source_paragraph = next((p for p in doc.paragraphs if p._element is el), None)
                 if source_paragraph is None: logger_fn.error(f"Пропуск параграфа {para_idx+1}: объект не найден."); para_idx += 1; continue
                 new_paragraph = result_doc.add_paragraph()
                 try:
                     # Обработка пустого параграфа
                     if not source_paragraph.text or source_paragraph.text.isspace():
                          if source_paragraph.style and hasattr(source_paragraph.style, 'name') and source_paragraph.style.name in result_doc.styles:
                              try: new_paragraph.style = result_doc.styles[source_paragraph.style.name]
                              except Exception: pass
                          if hasattr(source_paragraph, 'paragraph_format'): _safe_copy_paragraph_format(source_paragraph.paragraph_format, new_paragraph.paragraph_format)
                     # Обработка непустого параграфа
                     else:
                         tokens = tokenize_paragraph_universal(source_paragraph)
                         if not tokens: # Если нет токенов, копируем как есть
                             new_paragraph.text = source_paragraph.text
                             if source_paragraph.style and hasattr(source_paragraph.style, 'name') and source_paragraph.style.name in result_doc.styles:
                                 try: new_paragraph.style = result_doc.styles[source_paragraph.style.name]
                                 except Exception: pass
                             if hasattr(source_paragraph, 'paragraph_format'): _safe_copy_paragraph_format(source_paragraph.paragraph_format, new_paragraph.paragraph_format)
                         else: # Есть токены - ищем и реконструируем
                             matches = _find_matches_for_highlighting(tokens, search_lemmas_map, search_stems_map, search_phrase_lemmas_map, search_phrase_stems_map)
                             _reconstruct_paragraph_with_highlighting(new_paragraph, source_paragraph, tokens, matches, word_stats, phrase_stats) # Параметры сносок убраны
                 except Exception as e_para: # Обработка ошибки параграфа
                     logger_fn.error(f"Ошибка при обработке параграфа {para_idx+1}: {e_para}", exc_info=True)
                     # Fallback: добавляем как простой текст
                     try:
                         p_elem = new_paragraph._p; p_elem.getparent().remove(p_elem); fallback_p = result_doc.add_paragraph(source_paragraph.text if source_paragraph else "[Ошибка копирования параграфа]")
                         logger_fn.warning(f"Параграф {para_idx+1} добавлен как простой текст.")
                         if source_paragraph: # Попытка скопировать стиль/формат
                             if source_paragraph.style and hasattr(source_paragraph.style, 'name') and source_paragraph.style.name in result_doc.styles:
                                 try: fallback_p.style = result_doc.styles[source_paragraph.style.name]
                                 except Exception: pass
                             if hasattr(source_paragraph, 'paragraph_format'): _safe_copy_paragraph_format(source_paragraph.paragraph_format, fallback_p.paragraph_format)
                     except Exception as e_fallback: logger_fn.error(f"    Не удалось выполнить fallback для параграфа {para_idx+1}: {e_fallback}")
                 para_idx += 1

             # --- Обработка Таблицы ---
             elif el_type == 'table':
                 source_table = next((t for t in doc.tables if t._element is el), None)
                 if source_table is None: logger_fn.error(f"Пропуск таблицы {table_idx+1}: объект не найден."); table_idx += 1; continue
                 logger_fn.debug(f"  Обработка таблицы {table_idx+1}...")
                 try:
                     # Создаем новую таблицу
                     new_table = result_doc.add_table(rows=len(source_table.rows), cols=len(source_table.columns))
                     # Копируем стиль и свойства таблицы
                     if source_table.style and hasattr(source_table.style, 'name'):
                         try:
                             if source_table.style.name in result_doc.styles: new_table.style = result_doc.styles[source_table.style.name]
                         except Exception as e_tbl_style: logger_fn.warning(f"    Не удалось применить стиль таблицы: {e_tbl_style}")
                     try: new_table.autofit = source_table.autofit
                     except AttributeError: pass
                     try: new_table.alignment = source_table.alignment
                     except AttributeError: pass

                     # Обработка ячеек
                     for i, row in enumerate(source_table.rows):
                         if i >= len(new_table.rows): continue
                         new_row = new_table.rows[i]
                         for j, cell in enumerate(row.cells):
                             if j >= len(new_row.cells): continue
                             source_cell = cell; target_cell = new_table.cell(i, j); target_cell.text = "" # Очищаем
                             try: target_cell.vertical_alignment = source_cell.vertical_alignment
                             except Exception: pass
                             # Обработка параграфов внутри ячейки
                             for cell_para_idx, cell_paragraph in enumerate(source_cell.paragraphs):
                                 new_cell_paragraph = target_cell.add_paragraph()
                                 try:
                                     if not cell_paragraph.text or cell_paragraph.text.isspace(): # Пустой
                                          if cell_paragraph.style and hasattr(cell_paragraph.style,'name') and cell_paragraph.style.name in result_doc.styles:
                                             try: new_cell_paragraph.style = result_doc.styles[cell_paragraph.style.name]
                                             except Exception: pass
                                          if hasattr(cell_paragraph, 'paragraph_format'): _safe_copy_paragraph_format(cell_paragraph.paragraph_format, new_cell_paragraph.paragraph_format)
                                     else: # Непустой
                                         tokens = tokenize_paragraph_universal(cell_paragraph)
                                         if not tokens: # Нет токенов
                                             new_cell_paragraph.text = cell_paragraph.text
                                             if cell_paragraph.style and hasattr(cell_paragraph.style,'name') and cell_paragraph.style.name in result_doc.styles:
                                                try: new_cell_paragraph.style = result_doc.styles[cell_paragraph.style.name]
                                                except Exception: pass
                                             if hasattr(cell_paragraph, 'paragraph_format'): _safe_copy_paragraph_format(cell_paragraph.paragraph_format, new_cell_paragraph.paragraph_format)
                                         else: # Есть токены
                                             matches = _find_matches_for_highlighting(tokens, search_lemmas_map, search_stems_map, search_phrase_lemmas_map, search_phrase_stems_map)
                                             _reconstruct_paragraph_with_highlighting(new_cell_paragraph, cell_paragraph, tokens, matches, word_stats, phrase_stats) # Параметры сносок убраны
                                 except Exception as e_cp: # Ошибка параграфа ячейки
                                     logger_fn.error(f"Ошибка параграфа {cell_para_idx+1} в ячейке ({i},{j}): {e_cp}")
                                     # Fallback для параграфа ячейки
                                     try:
                                         p_elem = new_cell_paragraph._p; p_elem.getparent().remove(p_elem); fallback_cp = target_cell.add_paragraph(cell_paragraph.text)
                                         logger_fn.warning(f"      Параграф {cell_para_idx+1} в ячейке ({i},{j}) добавлен как текст.")
                                         if cell_paragraph.style and hasattr(cell_paragraph.style,'name') and cell_paragraph.style.name in result_doc.styles:
                                            try: fallback_cp.style = result_doc.styles[cell_paragraph.style.name]
                                            except Exception: pass
                                         if hasattr(cell_paragraph, 'paragraph_format'): _safe_copy_paragraph_format(cell_paragraph.paragraph_format, fallback_cp.paragraph_format)
                                     except Exception as e_fallback_cell: logger_fn.error(f"        Fallback ячейки ({i},{j}) не удался: {e_fallback_cell}")
                             # Удаление пустого первого параграфа
                             if len(target_cell.paragraphs) > 1 and not target_cell.paragraphs[0].text and not target_cell.paragraphs[0].runs:
                                 try: target_cell._element.remove(target_cell.paragraphs[0]._element)
                                 except Exception as e_remove_empty: logger_fn.warning(f"      Не удалось удалить пустой параграф ячейки ({i},{j}): {e_remove_empty}")
                 except Exception as e_tbl: # Ошибка таблицы
                     logger_fn.error(f"Критическая ошибка обработки таблицы {table_idx+1}: {e_tbl}", exc_info=True)
                     result_doc.add_paragraph(f"[Ошибка копирования таблицы {table_idx+1}]") # Добавляем заглушку
                 table_idx += 1

             else: # Неизвестный тип элемента
                 logger_fn.warning(f"Пропущен неизвестный тип элемента: {el.tag}")

        # --- Сохранение результата и финальная статистика ---
        logger_fn.info(f"Сохранение результата в {output_path}...")
        result_doc.save(output_path)
        end_process_time = time.time(); log_cache_stats(); 
        processing_time = end_process_time - start_process_time
        logger_fn.info(f"Документ обработан и сохранен за {processing_time:.2f} сек.")

        # Готовим статистику (вся найденная)
        final_results = {
            'word_stats': dict(word_stats),     # Преобразуем defaultdict в dict
            'phrase_stats': dict(phrase_stats) # Преобразуем defaultdict в dict
        }
        logger_fn.debug(f"Финальная статистика: {len(final_results['word_stats'])} слов, {len(final_results['phrase_stats'])} фраз.")
        return final_results

    # Обработка ошибок верхнего уровня
    except FileNotFoundError as e:
        logger_fn.error(f"Исходный файл не найден: {e}", exc_info=True)
        log_cache_stats();  # Логируем кеш перед выходом
        return None
    except Exception as e:
        logger_fn.error(f"Критическая ОШИБКА в analyze_and_highlight_terms_docx: {e}", exc_info=True) # Обновлено имя функции
        log_cache_stats(); 
        return None

# --- END OF FILE footnotes_service.py ---