import re
from itertools import product
from pathlib import Path

from services.fulltext_search.utils import normalize_text

try:
    from docx import Document
except ImportError:
    Document = None


CASES = ("nom", "gen", "dat", "acc", "ins", "pre")
VOWELS = set("аеёиоуыэюя")

INDECLINABLE_SUFFIXES = (
    "енко", "ко",
    "ых", "их",
    "аго", "яго", "ово",
    "е", "и", "о", "у", "ы", "э", "ю"
)

PARTICLES = {
    "де", "фон", "ван", "дер", "да", "дель", "ди", "ду", "ле", "ла"
}

EXCEPTION_FORMS = {
    "гайдай": {
        "male": {
            "nom": ["Гайдай"],
            "gen": ["Гайдая"],
            "dat": ["Гайдаю"],
            "acc": ["Гайдая"],
            "ins": ["Гайдаем"],
            "pre": ["Гайдае"],
        }
    },
    "дюма": {
        "male": {c: ["Дюма"] for c in CASES},
        "female": {c: ["Дюма"] for c in CASES},
    },
    "золя": {
        "male": {c: ["Золя"] for c in CASES},
        "female": {c: ["Золя"] for c in CASES},
    },
    "гавальда": {
        "male": {c: ["Гавальда"] for c in CASES},
        "female": {c: ["Гавальда"] for c in CASES},
    },
    "ли": {
        "male": {c: ["Ли"] for c in CASES},
        "female": {c: ["Ли"] for c in CASES},
    },
    "шоу": {
        "male": {c: ["Шоу"] for c in CASES},
        "female": {c: ["Шоу"] for c in CASES},
    },
    "гюго": {
        "male": {c: ["Гюго"] for c in CASES},
        "female": {c: ["Гюго"] for c in CASES},
    },
    "заяц": {
        "male": {
            "nom": ["Заяц"],
            "gen": ["Заяца", "Зайца"],
            "dat": ["Заяцу", "Зайцу"],
            "acc": ["Заяца", "Зайца"],
            "ins": ["Заяцем", "Зайцем"],
            "pre": ["Заяце", "Зайце"],
        }
    },
    "топчий": {
        "male": {
            "nom": ["Топчий"],
            "gen": ["Топчего", "Топчия"],
            "dat": ["Топчему", "Топчию"],
            "acc": ["Топчего", "Топчия"],
            "ins": ["Топчим", "Топчием"],
            "pre": ["Топчем", "Топчие"],
        }
    },
}


def uniq(seq):
    out = []
    seen = set()
    for x in seq:
        k = normalize_text(x)
        if k not in seen:
            seen.add(k)
            out.append(x)
    return out


def constant_forms(surname: str):
    return {case: [surname] for case in CASES}


def mk_forms(nom, gen, dat, acc, ins, pre):
    return {
        "nom": [nom],
        "gen": uniq(gen if isinstance(gen, list) else [gen]),
        "dat": uniq(dat if isinstance(dat, list) else [dat]),
        "acc": uniq(acc if isinstance(acc, list) else [acc]),
        "ins": uniq(ins if isinstance(ins, list) else [ins]),
        "pre": uniq(pre if isinstance(pre, list) else [pre]),
    }


def merge_forms(*forms_dicts):
    out = {case: [] for case in CASES}
    for d in forms_dicts:
        for case in CASES:
            out[case].extend(d[case])
    for case in CASES:
        out[case] = uniq(out[case])
    return out


def ends_with_consonant(s: str) -> bool:
    return bool(s) and s[-1] in "бвгджзклмнпрстфхцчшщ"


def is_particle(word: str) -> bool:
    return normalize_text(word) in PARTICLES


def is_safe_masc_adj(s: str) -> bool:
    return s.endswith(("ский", "цкий", "ской", "цкой", "ый"))


def is_fem_adj(s: str) -> bool:
    return s.endswith(("ская", "цкая", "ая", "яя"))


def adj_masc_forms(surname: str):
    s = normalize_text(surname)

    if s.endswith(("ский", "цкий", "ской", "цкой")):
        stem = surname[:-2]
        return mk_forms(
            surname,
            stem + "ого",
            stem + "ому",
            stem + "ого",
            stem + "им",
            stem + "ом",
        )

    if s.endswith("ий"):
        stem = surname[:-2]
        return mk_forms(
            surname,
            stem + "его",
            stem + "ему",
            stem + "его",
            stem + "им",
            stem + "ем",
        )

    if s.endswith(("ый", "ой")):
        stem = surname[:-2]
        return mk_forms(
            surname,
            stem + "ого",
            stem + "ому",
            stem + "ого",
            stem + "ым",
            stem + "ом",
        )

    return constant_forms(surname)


def adj_fem_forms(surname: str):
    s = normalize_text(surname)

    if s.endswith("яя"):
        stem = surname[:-2]
        return mk_forms(
            surname,
            stem + "ей",
            stem + "ей",
            stem + "юю",
            stem + "ей",
            stem + "ей",
        )

    stem = surname[:-2]
    return mk_forms(
        surname,
        stem + "ой",
        stem + "ой",
        stem + "ую",
        stem + "ой",
        stem + "ой",
    )


def noun_soft_sign_forms(surname: str):
    stem = surname[:-1]
    return mk_forms(
        surname,
        stem + "я",
        stem + "ю",
        stem + "я",
        stem + "ем",
        stem + "е",
    )


def noun_y_forms(surname: str):
    stem = surname[:-1]
    return mk_forms(
        surname,
        stem + "я",
        stem + "ю",
        stem + "я",
        stem + "ем",
        stem + "е",
    )


def noun_a_forms(surname: str):
    s = normalize_text(surname)
    stem = surname[:-1]
    prev = s[-2] if len(s) >= 2 else ""
    gen_end = "и" if prev in "гкхжчшщц" else "ы"
    return mk_forms(
        surname,
        stem + gen_end,
        stem + "е",
        stem + "у",
        stem + "ой",
        stem + "е",
    )


def noun_ya_forms(surname: str):
    stem = surname[:-1]
    return mk_forms(
        surname,
        stem + "и",
        stem + "е",
        stem + "ю",
        stem + "ей",
        stem + "е",
    )


def masc_consonant_forms(surname: str, foreign_ov_in: bool = False):
    s = normalize_text(surname)
    ins = surname + "ом"

    if s.endswith(("ов", "ев", "ёв", "ин", "ын")) and not foreign_ov_in:
        ins = surname + "ым"

    return mk_forms(
        surname,
        surname + "а",
        surname + "у",
        surname + "а",
        ins,
        surname + "е",
    )


def compound_space_forms(surname: str, gender: str, meta: dict):
    words = surname.split()
    if len(words) >= 2 and is_particle(words[0]):
        first = words[0]
        tail = " ".join(words[1:])
        tail_forms = decline_surname_forms(tail, gender, meta)
        return {case: [first + " " + x for x in tail_forms[case]] for case in CASES}
    return None


def compound_hyphen_forms(surname: str, gender: str, meta: dict):
    if "-" not in surname:
        return None

    parts = surname.split("-")
    part_metas = meta.get("hyphen_parts") or [{} for _ in parts]
    if len(part_metas) != len(parts):
        part_metas = [{} for _ in parts]

    declined_parts = []
    for part, part_meta in zip(parts, part_metas):
        if is_particle(part):
            declined_parts.append(constant_forms(part))
        else:
            declined_parts.append(decline_surname_forms(part, gender, part_meta))

    out = {case: [] for case in CASES}
    for case in CASES:
        variants = [d[case] for d in declined_parts]
        out[case] = uniq(["-".join(x) for x in product(*variants)])
    return out


def decline_surname_forms(surname: str, gender: str, meta: dict | None = None):
    meta = meta or {}
    gender = gender.lower().strip()
    s = normalize_text(surname)

    if s in EXCEPTION_FORMS and gender in EXCEPTION_FORMS[s]:
        return EXCEPTION_FORMS[s][gender]

    spaced = compound_space_forms(surname, gender, meta)
    if spaced:
        return spaced

    hyphenated = compound_hyphen_forms(surname, gender, meta)
    if hyphenated:
        return hyphenated

    override = meta.get("override_forms")
    if override:
        return override

    if s.endswith("иа") and not s.endswith("ия"):
        return constant_forms(surname)

    if len(s) >= 2 and s.endswith("а") and s[-2] in VOWELS:
        return constant_forms(surname)

    for suf in INDECLINABLE_SUFFIXES:
        if s.endswith(suf):
            return constant_forms(surname)

    if meta.get("indeclinable_final_a"):
        return constant_forms(surname)

    if gender == "female" and is_fem_adj(s):
        return adj_fem_forms(surname)

    if gender == "male":
        if is_safe_masc_adj(s):
            return adj_masc_forms(surname)

        if s.endswith(("ий", "ой")):
            mode = meta.get("y_declension", "auto")
            if mode == "adjective":
                return adj_masc_forms(surname)
            if mode == "noun":
                return noun_y_forms(surname)
            if mode == "both":
                return merge_forms(adj_masc_forms(surname), noun_y_forms(surname))
            return adj_masc_forms(surname)

    if gender == "female" and (ends_with_consonant(s) or s.endswith(("ь", "й"))):
        return constant_forms(surname)

    if gender == "male" and s.endswith("ь"):
        return noun_soft_sign_forms(surname)

    if gender == "male" and s.endswith("й"):
        mode = meta.get("y_declension", "noun")
        if mode == "adjective":
            return adj_masc_forms(surname)
        if mode == "both":
            return merge_forms(adj_masc_forms(surname), noun_y_forms(surname))
        return noun_y_forms(surname)

    if s.endswith("я"):
        return noun_ya_forms(surname)

    if s.endswith("а"):
        return noun_a_forms(surname)

    if gender == "male" and ends_with_consonant(s):
        return masc_consonant_forms(
            surname,
            foreign_ov_in=bool(meta.get("foreign_ov_in"))
        )

    return constant_forms(surname)


def build_surname_index(surnames: list[dict]) -> dict:
    index = {}

    for item in surnames:
        base = item["surname"].strip()
        gender = item["gender"].strip().lower()
        meta = item.get("meta", {})

        forms = decline_surname_forms(base, gender, meta)

        for case in CASES:
            for form in forms[case]:
                key = normalize_text(form)
                index.setdefault(key, []).append({
                    "base_surname": base,
                    "gender": gender,
                    "matched_case": case,
                    "meta": meta,
                })

    return index


def tokenize_text(text: str):
    pattern = re.compile(r"[А-Яа-яЁёA-Za-z]+(?:-[А-Яа-яЁёA-Za-z]+)?")
    return [(m.group(0), m.start(), m.end()) for m in pattern.finditer(text)]


def find_surnames_in_text(text: str, surname_index: dict):
    results = []
    for token, start, end in tokenize_text(text):
        key = normalize_text(token)
        if key in surname_index:
            for hit in surname_index[key]:
                results.append({
                    "found": token,
                    "normalized": key,
                    "base_surname": hit["base_surname"],
                    "gender": hit["gender"],
                    "matched_case": hit["matched_case"],
                    "start": start,
                    "end": end,
                    "context": text[max(0, start - 40):min(len(text), end + 40)],
                })
    return results


def read_txt(path: str) -> str:
    return Path(path).read_text(encoding="utf-8")


def read_docx(path: str) -> str:
    if Document is None:
        raise ImportError("Install python-docx for .docx")
    doc = Document(path)
    parts = [p.text for p in doc.paragraphs if p.text]
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                if cell.text:
                    parts.append(cell.text)
    return "\n".join(parts)


def read_document(path: str) -> str:
    ext = Path(path).suffix.lower()
    if ext == ".txt":
        return read_txt(path)
    if ext == ".docx":
        return read_docx(path)
    raise ValueError(f"Unsupported format: {ext}")


def find_surnames_in_document(path: str, surnames: list[dict]):
    text = read_document(path)
    index = build_surname_index(surnames)
    return find_surnames_in_text(text, index)


if __name__ == "__main__":
    surnames = [
        {"surname": "Иванов", "gender": "male"},
        {"surname": "Дарвин", "gender": "male", "meta": {"foreign_ov_in": True}},
        {"surname": "Блок", "gender": "female"},
        {"surname": "Берия", "gender": "male"},
        {"surname": "Галуа", "gender": "male"},
        {"surname": "Золя", "gender": "male", "meta": {"indeclinable_final_a": True}},
        {"surname": "Топчий", "gender": "male", "meta": {"y_declension": "both"}},
        {"surname": "де Голль", "gender": "male"},
        {"surname": "Скворцов-Степанов", "gender": "male"},
    ]

    sample_text = """
    Договор подписан Ивановым.
    Теория предложена Дарвином.
    Письмо отправлено Анне Блок.
    Материалы о Берии приобщены к делу.
    Переписка с Галуа найдена в архиве.
    Речь шла об Эмиле Золя.
    Справка выдана Топчию.
    Протокол составлен де Голлем и Скворцовым-Степановым.
    """

    index = build_surname_index(surnames)
    matches = find_surnames_in_text(sample_text, index)

    for m in matches:
        print(
            f"{m['found']} -> {m['base_surname']} "
            f"({m['gender']}, {m['matched_case']})"
        )
