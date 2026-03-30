import docx
import time
from dataclasses import dataclass
from typing import List, Union, Optional, Tuple, NamedTuple
from functools import cmp_to_key
from docx.text.paragraph import Paragraph
from docx.text.run import Run
from docx.text.hyperlink import Hyperlink
from docx.table import Table
from docx.oxml.ns import qn
from docx.oxml import OxmlElement, CT_R, CT_Text, CT_P, CT_Hyperlink
from copy import deepcopy

from services.progress.combined_progress.combined_progress import CombinedProgress
from services.progress.combined_progress.process_particle import ProgressParticle
from services.utils.intersects_at import intersects_at
from services.utils.interval import Interval
from services.analysis.analyser import Analyser
from services.analysis.analysis_match import AnalysisMatch
from services.analysis.annot_content import get_annot_title_content, get_multiple_get_annot_title_content
from services.fulltext_search.fulltext_search import FulltextSearch, SearchStrategy
from services.tokenization import Token, TokenType, TokenDictionary, Tokenizer
from services.fulltext_search.phrase import Phrase
from services.utils.timeit import timeit


class _SearchIntersection(NamedTuple):
    start: int
    end: int
    r_start: int
    r_end: int


@dataclass
class _RunMapEntry:
    start: int
    end: int
    text: str
    run: CT_R
    search_intersection: Union[_SearchIntersection, None] = None


class AnalyserDocx(Analyser):
    _PARTICLE_PREPARATION: str = 'docx_preparation'
    _PARTICLE_SEARCH: str = 'docx_search'
    _DOCX_PROGRESS_EMIT_EVERY: int = 100

    document: docx.Document
    _tokenize_time_total: float
    _global_document_dictionary: Optional[TokenDictionary]
    _search_phrases: List[Phrase]
    _progress: Optional[CombinedProgress]
    _docx_preparation_value: float
    _docx_search_value: float

    def __init__(self, document: Union[docx.Document, str]):
        super().__init__()
        if isinstance(document, str):
            document = docx.Document(document)

        self.document = document
        self._tokenize_time_total = 0.0
        self._global_document_dictionary = None
        self._search_phrases = []
        self._progress = None
        self._docx_preparation_value = 0.0
        self._docx_search_value = 0.0

    def _docx_progress_preparation_value(self, value: float) -> None:
        self._docx_preparation_value = value

        if self._progress is None:
            return

        if int(value) % AnalyserDocx._DOCX_PROGRESS_EMIT_EVERY != 0:
            return

        self._progress.set_particle_value(
            AnalyserDocx._PARTICLE_PREPARATION,
            value,
        )

    def _docx_progress_search_value(self, value: float) -> None:
        self._docx_search_value = value

        if self._progress is None:
            return

        if int(value) % AnalyserDocx._DOCX_PROGRESS_EMIT_EVERY != 0:
            return

        self._progress.set_particle_value(
            AnalyserDocx._PARTICLE_SEARCH,
            value,
        )

    def _docx_progress_flush_preparation(self) -> None:
        if self._progress is None:
            return

        self._progress.set_particle_value(
            AnalyserDocx._PARTICLE_PREPARATION,
            self._docx_preparation_value,
        )

    def _docx_progress_flush_search(self) -> None:
        if self._progress is None:
            return

        self._progress.set_particle_value(
            AnalyserDocx._PARTICLE_SEARCH,
            self._docx_search_value,
        )

    @staticmethod
    def __clone_run(
            source_run: CT_R,
            new_text: str,
            highlight: bool,
            highlight_val: str,
            skip_break_lines: bool = False,
    ) -> CT_R:
        new_run: CT_R = OxmlElement('w:r')

        for child in source_run:
            if skip_break_lines and child.tag == qn('w:br'):
                continue

            new_run.append(deepcopy(child))

        # update text
        text_element: CT_Text = new_run.find(qn('w:t'))
        text_element.text = new_text
        text_element.set(qn('xml:space'), 'preserve')

        # highlight text
        if highlight:
            rPr = new_run.find(qn('w:rPr'))

            if rPr is None:
                rPr = OxmlElement('w:rPr')
                new_run.insert(0, rPr)

            shd = OxmlElement('w:shd')
            shd.set(qn('w:val'), 'clear')
            shd.set(qn('w:fill'), highlight_val)
            shd.set(qn('w:color'), 'auto')
            rPr.append(shd)

        return new_run

    @staticmethod
    def __isolate_new_run_xml(
            batch: List[CT_R],
            phrase_start_index: int,
            phrase_end_index: int,
            highlight_val: str
    ) -> Tuple[List[CT_R], Optional[CT_R]]:
        phrase_end_index -= 1

        # [start] build map
        run_map: List[_RunMapEntry] = []
        start_index = 0

        for run in batch:
            run_map.append(_RunMapEntry(
                start=start_index,
                end=start_index + len(run.text) - 1,
                text=run.text,
                run=run,
            ))
            start_index += len(run.text)
        # [end]

        # [start] get run with phrase
        for map_item in run_map:
            run_start_index = map_item.start
            run_end_index = map_item.end
            intersection = intersects_at((run_start_index, run_end_index), (phrase_start_index, phrase_end_index))

            if intersection:
                intersection_len = intersection[1] - intersection[0]
                r_start = intersection[0] - run_start_index
                r_end = r_start + intersection_len
                map_item.search_intersection = _SearchIntersection(
                    start=intersection[0],
                    end=intersection[1],
                    r_start=r_start,
                    r_end=r_end,
                )
        # [end]

        # [start] get run with phrase
        new_batch: List[CT_R] = []
        run_match_el: Optional[CT_R] = None

        for map_item in run_map:
            if map_item.search_intersection is None:
                new_batch.append(map_item.run)
                continue

            # [start] split run text on three new parts: before, match and after
            text = map_item.run.text
            i_start = map_item.search_intersection.r_start
            i_end = map_item.search_intersection.r_end
            part_before_match: str = text[:i_start]
            part_match: str = text[i_start:i_end + 1]
            part_after_match: str = text[i_end + 1:]

            if len(map_item.run.br_lst) == 1:
                part_before_match = part_before_match.lstrip("\n")
                skip_break_lines = False
            else:
                skip_break_lines = True

            run_before_match = AnalyserDocx.__clone_run(map_item.run, part_before_match, False, highlight_val,
                                                        skip_break_lines=skip_break_lines)
            run_match = AnalyserDocx.__clone_run(map_item.run, part_match, True, highlight_val, skip_break_lines=True)
            run_after_match = AnalyserDocx.__clone_run(map_item.run, part_after_match, False, highlight_val,
                                                       skip_break_lines=True)

            parent = map_item.run.getparent()
            index = list(parent).index(map_item.run)
            parent.remove(map_item.run)
            for i, run_el in enumerate([run_before_match, run_match, run_after_match]):
                parent.insert(index + i, run_el)

            new_batch.append(run_before_match)
            new_batch.append(run_match)
            new_batch.append(run_after_match)
            run_match_el = run_match
            # [end]
        # [end]

        return (new_batch, run_match_el)

    def __search_all_phrases(
            self,
            source_tokens: List[Token],
            search_phrases: List[Phrase]
    ) -> List[AnalysisMatch]:
        """Search all phrases using optimized strategy with dictionary."""
        fulltext_search = FulltextSearch(source_tokens)
        search_phrases_for_search: List[Tuple[Phrase, List[Token]]] = [
            (phrase, phrase.tokens) for phrase in search_phrases
        ]
        regex_patterns_dict = None

        if self.analyse_data.regex_patterns:
            regex_patterns_dict = {
                pattern.pattern_name: pattern
                for pattern in self.analyse_data.regex_patterns
            }

        phrase_results = fulltext_search.search_all(
            search_phrases=search_phrases_for_search,
            text_strategy=SearchStrategy.FUZZY_WORDS_PUNCT,
            search_patterns=regex_patterns_dict
        )

        # [start] todo: dev only simplify
        fts_matches = []

        for _, fts_match in phrase_results:
            fts_matches.extend(fts_match)
        # [end]

        matches = self._convert_fts_matches(fts_matches)

        # [start] sort matches by position in text
        def cmp(m1: AnalysisMatch, m2: AnalysisMatch):
            return m1.search_match.start_token_idx - m2.search_match.start_token_idx

        matches = sorted(matches, key=cmp_to_key(cmp))
        # [end]

        return matches

    def __process_batch(self, batch: List[CT_R], paragraph: Paragraph) -> None:
        # [start] find matches in concatenated batch text
        text = ''.join([run.text for run in batch])
        start_time = time.time()
        source_tokens: List[Token] = Tokenizer(None).tokenize_text(text)
        self._tokenize_time_total += time.time() - start_time

        matches: List[AnalysisMatch] = self.__search_all_phrases(
            source_tokens,
            self._search_phrases
        )
        # [end]

        # [start] highlight match in document
        match_run_els_list: List[Tuple[AnalysisMatch, List[CT_R]]] = []

        self._all_matches.extend(matches)

        for match in matches:
            highlight_val: str = self._highlight_color_for_match(match).rrggbb().upper()
            search_match = match.search_match
            start_token_idx = search_match.start_token_idx
            end_token_idx = search_match.end_token_idx

            # [start] fill match.runs
            match_run_els: List[CT_R] = []

            for i in range(start_token_idx, end_token_idx + 1):
                token = source_tokens[i]
                batch, run_match_el = self.__isolate_new_run_xml(batch, token.start, token.end, highlight_val)

                if run_match_el is not None:
                    match_run_els.append(run_match_el)

            if match.runs is None:
                match.runs = []

            match.runs.extend(match_run_els)
            # [end]
        # [end]

        # [start] build footnotes map
        footnotes_map: dict[Interval, List[AnalysisMatch]] = {}

        for match in matches:
            search_match = match.search_match
            start_token_idx = search_match.start_token_idx
            end_token_idx = search_match.end_token_idx
            new_match_interval = Interval(start_token_idx, end_token_idx)
            overlap_with: Optional[Interval] = None

            # find match overlap with other matches
            for interval, _ in footnotes_map.items():
                if new_match_interval.intersects(interval):
                    overlap_with = interval
                    break

            # place new match in map
            if overlap_with:
                # union match with others
                footnotes_map[overlap_with].append(match)
                interval_matches: List[AnalysisMatch] = footnotes_map[overlap_with]
                new_interval = overlap_with.union(new_match_interval)
                del footnotes_map[overlap_with]
                footnotes_map[new_interval] = interval_matches
            else:
                # new match not overlaps with other
                footnotes_map[new_match_interval] = [match]

        # create footnotes
        for _, matches in footnotes_map.items():
            if len(matches) == 1:
                title, content = get_annot_title_content(matches[0])
            else:
                title, content = get_multiple_get_annot_title_content(matches)

            runs: List[CT_R] = []

            for match in matches:
                runs.extend(match.runs)

            comment_runs: List[Run] = [Run(r_el, paragraph) for r_el in runs]

            self.document.add_comment(
                runs=comment_runs,
                text=content,
                author=title,
                initials="",
            )
        # [end]

        # [start] add annotation to highlight text in document
        for match, match_run_els in match_run_els_list:
            title, content = get_annot_title_content(match)
            comment_runs: List[Run] = [Run(r_el, paragraph) for r_el in match_run_els]

            self.document.add_comment(
                runs=comment_runs,
                text=content,
                author=title,
                initials="",
            )
        # [end]

    @staticmethod
    def __split_on_batches(element: Union[CT_P, CT_Hyperlink]) -> List[List[CT_R]]:
        batches = []
        batch_runs = []

        for child in element.r_lst:
            batch_runs.append(child)

        if len(batch_runs) > 0:
            batches.append(batch_runs)

        return batches

    def __analyse_paragraph(self, paragraph: Paragraph) -> None:
        batches = AnalyserDocx.__split_on_batches(paragraph._element)

        for batch in batches:
            self.__process_batch(batch, paragraph)

        hyperlinks: List[Hyperlink] = paragraph.hyperlinks

        for link in hyperlinks:
            self.__analyze_link(link._element, paragraph)

    def __analyze_link(self, link: CT_Hyperlink, paragraph: Paragraph) -> None:
        batches = AnalyserDocx.__split_on_batches(link)

        for batch in batches:
            self.__process_batch(batch, paragraph)

    def __build_global_dictionary(self) -> TokenDictionary:
        """Build dictionary from entire document text."""
        self._docx_progress_preparation_value(0.0)

        paragraphs: List[Paragraph] = self.document.paragraphs
        tables: List[Table] = self.document.tables

        # [start] tokenize and build dictionary gradually
        all_tokens: List[Token] = []
        preparation_paragraph_index: int = 0

        for paragraph in paragraphs:
            paragraph_tokens = Tokenizer(None).tokenize_text(paragraph.text)
            all_tokens.extend(paragraph_tokens)
            preparation_paragraph_index += 1
            self._docx_progress_preparation_value(float(preparation_paragraph_index))

        for table in tables:
            for row in table.rows:
                for cell in row.cells:
                    for paragraph in cell.paragraphs:
                        paragraph_tokens = Tokenizer(None).tokenize_text(paragraph.text)
                        all_tokens.extend(paragraph_tokens)
                        preparation_paragraph_index += 1
                        self._docx_progress_preparation_value(float(preparation_paragraph_index))
        # [end]

        dictionary = TokenDictionary(all_tokens)

        return dictionary

    def __filter_phrases_by_dictionary(
            self,
            search_phrases: List[Phrase],
            dictionary: TokenDictionary
    ) -> List[Phrase]:
        """Filter phrases: exclude those where at least one word is missing in dictionary."""
        filtered_phrases: List[Phrase] = []
        preparation_value: float = self._docx_preparation_value

        for phrase in search_phrases:
            search_words = [t for t in phrase.tokens if t.type == TokenType.WORD]

            preparation_value += 1.0

            if len(search_words) == 0:
                self._docx_progress_preparation_value(preparation_value)
                continue

            filtered_words = dictionary.filter_tokens(search_words)

            if len(filtered_words) == len(search_words):
                filtered_phrases.append(phrase)

            self._docx_progress_preparation_value(preparation_value)

        return filtered_phrases

    @timeit
    def analyse_and_highlight(self, task_id: Optional[str] = None) -> dict:
        self._all_matches: List[AnalysisMatch] = []
        self._search_phrases = []

        paragraphs: List[Paragraph] = self.document.paragraphs
        tables: List[Table] = self.document.tables
        phrases_list = self.analyse_data.phrases

        self._progress = None
        self._docx_preparation_value = 0.0
        self._docx_search_value = 0.0

        if task_id is not None:
            total_paragraphs_for_dict: int = len(paragraphs)
            total_table_paragraphs: int = 0

            for table in tables:
                for row in table.rows:
                    for cell in row.cells:
                        cell_paragraph_count: int = len(cell.paragraphs)
                        total_paragraphs_for_dict += cell_paragraph_count
                        total_table_paragraphs += cell_paragraph_count

            preparation_max: int = total_paragraphs_for_dict + len(phrases_list)
            search_max: int = len(paragraphs) + total_table_paragraphs

            self._progress = CombinedProgress(task_id, [
                ProgressParticle(
                    key=AnalyserDocx._PARTICLE_PREPARATION,
                    description='Подготовка',
                    max_value=preparation_max,
                ),
                ProgressParticle(
                    key=AnalyserDocx._PARTICLE_SEARCH,
                    description='Поиск и подсветка',
                    max_value=search_max,
                ),
            ])

        # [start] build global dictionary and filter search phrases by it
        self._global_document_dictionary = self.__build_global_dictionary()
        self._search_phrases = self.__filter_phrases_by_dictionary(
            phrases_list,
            self._global_document_dictionary
        )
        self._docx_progress_flush_preparation()
        # [end]

        # process paragraphs
        search_paragraph_index: int = 0

        for paragraph in paragraphs:
            self.__analyse_paragraph(paragraph)
            search_paragraph_index += 1
            self._docx_progress_search_value(float(search_paragraph_index))

        # process tables
        for table in tables:
            for row in table.rows:
                for cell in row.cells:
                    cell_paragraphs: List[Paragraph] = cell.paragraphs

                    for paragraph in cell_paragraphs:
                        self.__analyse_paragraph(paragraph)
                        search_paragraph_index += 1
                        self._docx_progress_search_value(float(search_paragraph_index))

        self._docx_progress_flush_search()

        return self._get_stats_result(self._all_matches)

    def save(self, output_path: str) -> None:
        self.document.save(output_path)
