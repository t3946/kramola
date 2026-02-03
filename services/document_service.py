# --- START OF FILE document_service.py ---

import os
import docx # Для работы с .docx файлами
import re   # Для регулярных выражений
import logging # Для логирования

# --- Настройка логирования ---
logger_ds = logging.getLogger(__name__)


# --- Константы ---
TOKEN_PATTERN = re.compile(r'\b\w+\b', re.UNICODE) # Паттерн для токенизации

# --- Функция для сохранения файлов ---
def save_uploaded_file(file_obj, upload_folder, custom_filename=None):
    """
    Сохраняет загруженный файловый объект в указанную папку.
    """
    if not file_obj or not file_obj.filename:
        logger_ds.error("Попытка сохранить пустой или некорректный file_obj.")
        return None
    filename = custom_filename if custom_filename else file_obj.filename
    filename = os.path.basename(filename) # Безопасность
    if not filename:
        logger_ds.error(f"Не удалось получить безопасное имя файла из: {custom_filename or file_obj.filename}")
        return None
    file_path = os.path.join(upload_folder, filename)
    try:
        os.makedirs(upload_folder, exist_ok=True)
        # Assuming file_obj is a Flask/Werkzeug FileStorage object
        if hasattr(file_obj, 'save') and callable(file_obj.save):
             file_obj.save(file_path)
        else:
             # Fallback for other file-like objects if needed
             with open(file_path, 'wb') as f:
                 # Read in chunks to handle large files
                 chunk_size = 8192
                 while True:
                     chunk = file_obj.read(chunk_size)
                     if not chunk:
                         break
                     f.write(chunk)
        logger_ds.info(f"Файл '{getattr(file_obj, 'filename', 'N/A')}' сохранен как '{file_path}'")
        return file_path
    except Exception as e:
        logger_ds.error(f"Ошибка сохранения файла '{getattr(file_obj, 'filename', 'N/A')}' в '{file_path}': {e}", exc_info=True)
        return None

# --- Функции для DOCX ---
def extract_words_from_docx(file_path, as_text=False):
    """
    Извлекает все слова (токены) или полный текст из файла DOCX.
    """
    logger_ds.debug(f"Извлечение {'текста' if as_text else 'слов'} из DOCX: {file_path}")
    if not os.path.exists(file_path):
        logger_ds.error(f"Файл DOCX не найден для извлечения: {file_path}")
        raise FileNotFoundError(f"Файл DOCX не найден: {file_path}")
    try:
        doc = docx.Document(file_path)
    except Exception as e:
        logger_ds.error(f"Ошибка открытия DOCX документа '{file_path}': {e}", exc_info=True)
        raise IOError(f"Ошибка при открытии документа {file_path}: {str(e)}")

    words = []
    full_text_parts = []
    try:
        for p in doc.paragraphs:
            text = p.text
            if text:
                full_text_parts.append(text)
                if not as_text:
                    words.extend(TOKEN_PATTERN.findall(text))
        for t in doc.tables:
            for r in t.rows:
                for c in r.cells:
                    for p_cell in c.paragraphs:
                        text = p_cell.text
                        if text:
                            full_text_parts.append(text)
                            if not as_text:
                                words.extend(TOKEN_PATTERN.findall(text))
    except Exception as e:
        logger_ds.error(f"Ошибка во время обработки содержимого DOCX '{file_path}': {e}", exc_info=True)
        pass # Continue processing if possible, or re-raise if critical
    if as_text:
        result = '\n'.join(full_text_parts)
        logger_ds.debug(f"Извлечен текст DOCX длиной {len(result)} символов.")
        return result
    else:
        logger_ds.debug(f"Извлечено {len(words)} слов из DOCX.")
        return words

def extract_lines_from_docx(file_path):
    """
    Извлекает непустые строки текста из файла DOCX.
    """
    logger_ds.debug(f"Извлечение строк из DOCX: {file_path}")
    if not os.path.exists(file_path):
        logger_ds.error(f"Файл DOCX не найден для извлечения строк: {file_path}")
        raise FileNotFoundError(f"Файл DOCX не найден: {file_path}")
    lines = []
    try:
        doc = docx.Document(file_path)
        for paragraph in doc.paragraphs:
            text = paragraph.text.strip()
            if text: lines.append(text)
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for paragraph in cell.paragraphs:
                        text = paragraph.text.strip()
                        if text: lines.append(text)
        logger_ds.debug(f"Извлечено {len(lines)} строк из DOCX.")
        return lines
    except Exception as e:
        logger_ds.error(f"Ошибка при извлечении строк из DOCX '{file_path}': {e}", exc_info=True)
        return []

# --- END OF FILE document_service.py ---