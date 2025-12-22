import docx
from typing import List, Union, Optional
from collections import defaultdict, Counter

from docx.text.paragraph import Paragraph
from docx.text.hyperlink import Hyperlink
from docx.table import Table
from docx.oxml.ns import qn
from docx.oxml import OxmlElement, CT_R, CT_Text, CT_P, CT_Hyperlink

from services.analyser import AnalyseData
from services.analyser.fulltext_search import FulltextSearch
from services.highlight_service import _find_matches_in_paragraph_tokens
from services.analyser.fulltext_search import Match
from utils.timeit import timeit
from copy import deepcopy
from services.task.progress import Progress


class AnalyserDocx:
    document: docx.Document
    analyse_data: AnalyseData
    word_stats: defaultdict
    phrase_stats: defaultdict

    def __init__(self, document: Union[docx.Document, str]):
        if isinstance(document, str):
            document = docx.Document(document)

        self.document = document
        self.word_stats = defaultdict(lambda: {'count': 0, 'forms': Counter()})
        self.phrase_stats = defaultdict(lambda: {'count': 0, 'forms': Counter()})

    def set_analyse_data(self, analyse_data: AnalyseData) -> None:
        self.analyse_data = analyse_data

    @staticmethod
    def __clone_run(source_run: CT_R, new_text: str, highlight=False) -> CT_R:
        new_run: CT_R = OxmlElement('w:r')

        for child in source_run:
            new_run.append(deepcopy(child))

        # update text
        text_element: CT_Text = new_run.find(qn('w:t'))
        text_element.text = new_text
        text_element.set(qn('xml:space'), 'preserve')

        # update highlight
        if highlight:
            rPr = new_run.find(qn('w:rPr'))

            if rPr is None:
                rPr = OxmlElement('w:rPr')
                new_run.insert(0, rPr)

            highlight = OxmlElement('w:highlight')
            highlight.set(qn('w:val'), 'green')
            rPr.append(highlight)

        return new_run

    @staticmethod
    def __isolate_new_run_xml(batch: List[CT_R], phrase_start_index: int, phrase_end_index: int) -> List[CT_R]:
        # after Run split operation, batch need to be updated
        new_batch = []

        # [start] find run by char index
        proceed_chars = 0

        for run_source in batch:
            source_run_text_element: CT_Text = run_source.find(qn('w:t'))

            if source_run_text_element is None:
                new_batch.append(run_source)
                continue

            text: str = source_run_text_element.text
            run_start_index = proceed_chars
            run_end_index = run_start_index + len(text)

            if run_start_index <= phrase_start_index <= run_end_index:
                match_length: int = phrase_end_index - phrase_start_index
                run_relative_char_start_index: int = phrase_start_index - proceed_chars
                run_relative_char_end_index: int = run_relative_char_start_index + match_length

                # [start] split run text on three new parts: before, match and after
                part_before_match: str = text[:run_relative_char_start_index]
                part_match: str = text[run_relative_char_start_index:run_relative_char_end_index]
                part_after_match: str = text[run_relative_char_end_index:]

                # needs to avoid cases "This is an apple" -> "This is anapple"
                source_run_text_element.set(qn('xml:space'), 'preserve')

                run_match = AnalyserDocx.__clone_run(run_source, part_match, True)
                run_after = AnalyserDocx.__clone_run(run_source, part_after_match)
                source_run_text_element.text = part_before_match
                run_source.addnext(run_match)
                run_match.addnext(run_after)
                # [end]

                new_batch.append(run_source)
                new_batch.append(run_match)
                new_batch.append(run_after)
            else:
                new_batch.append(run_source)

            proceed_chars += len(text)
        # [end]

        return new_batch

    def __update_match_statistics(self, match: dict, tokens: list) -> None:
        start_token_idx = match['start_token_idx']
        end_token_idx = match['end_token_idx']

        lemma_key = match['lemma_key']

        if match['type'] == 'phrase':
            phrase_key_str = " ".join(lemma_key)
            stats = self.phrase_stats[phrase_key_str]
            text_parts = []

            for i in range(start_token_idx, end_token_idx + 1):
                if i < len(tokens):
                    text_parts.append(tokens[i]['text'])

            found_text = "".join(text_parts).strip()
            stats['count'] += 1
            stats['forms'][found_text] += 1
        elif match['type'] == 'word':
            if lemma_key and len(lemma_key) == 1:
                word_lemma = lemma_key[0]
                stats = self.word_stats[word_lemma]

                if start_token_idx < len(tokens):
                    found_text = tokens[start_token_idx]['text']
                    stats['count'] += 1
                    stats['forms'][found_text.lower()] += 1

    def __process_batch(self, batch: List[CT_R]) -> None:
        text = ''

        # [start] find matches in concatenated batch text
        for run in batch:
            text += run.text

        source_tokens = FulltextSearch.tokenize_text(text)
        phrase_map = self.analyse_data.phrase_map if self.analyse_data.phrase_map else {}
        matches: List[Match] = _find_matches_in_paragraph_tokens(
            source_tokens,
            self.analyse_data.lemmas,
            self.analyse_data.stems,
            phrase_map
        )
        # [end]

        for match in matches:
            start_token_idx = match['start_token_idx']
            end_token_idx = match['end_token_idx']

            self.__update_match_statistics(match, source_tokens)

            for i in range(start_token_idx, end_token_idx + 1):
                token = source_tokens[i]
                batch = self.__isolate_new_run_xml(batch, token['start'], token['end'])

    @staticmethod
    def __split_on_batches(element: Union[CT_P, CT_Hyperlink]) -> List[List[CT_R]]:
        batches = []
        batch_runs = []
        qn_ct_r = qn('w:r')

        # [start] split paragraph text on batches of runs
        for child in element:
            if child.tag == qn_ct_r:
                batch_runs.append(child)
            elif len(batch_runs) > 0:
                batches.append(batch_runs)
                batch_runs = []

        if len(batch_runs) > 0:
            batches.append(batch_runs)
        # [end]

        return batches

    def __analyse_paragraph(self, paragraph: Paragraph) -> None:
        batches = AnalyserDocx.__split_on_batches(paragraph._element)

        for batch in batches:
            self.__process_batch(batch)

        hyperlinks: List[Hyperlink] = paragraph.hyperlinks

        for link in hyperlinks:
            self.__analyze_link(link._element)

    def __analyze_link(self, link: CT_Hyperlink) -> None:
        batches = AnalyserDocx.__split_on_batches(link)

        for batch in batches:
            self.__process_batch(batch)

    @timeit
    def analyse_and_highlight(self, task_id: Optional[str] = None) -> dict:
        #[start] reset stats
        self.word_stats = defaultdict(lambda: {'count': 0, 'forms': Counter()})
        self.phrase_stats = defaultdict(lambda: {'count': 0, 'forms': Counter()})
        #[end]

        paragraphs: List[Paragraph] = self.document.paragraphs
        tables: List[Table] = self.document.tables

        # [start] count progress max
        progress = None

        if task_id is not None:
            total_table_paragraphs = 0

            for table in tables:
                for row in table.rows:
                    for cell in row.cells:
                        total_table_paragraphs += len(cell.paragraphs)

            total_items = len(paragraphs) + total_table_paragraphs
            progress = Progress(task_id, max_value=total_items)
        # [end]

        # process paragraphs
        for paragraph in paragraphs:
            self.__analyse_paragraph(paragraph)

            if progress:
                progress.add(1)

        # process tables
        for table in tables:
            for row in table.rows:
                for cell in row.cells:
                    cell_paragraphs: List[Paragraph] = cell.paragraphs

                    for paragraph in cell_paragraphs:
                        self.__analyse_paragraph(paragraph)

                        if progress:
                            progress.add(1)

        if progress:
            progress.flush()
            progress.clear()

        #[start] return stats in same format as analyze_and_highlight_pdf
        final_ws = {l: {'c': d['count'], 'f': dict(d['forms'])} for l, d in self.word_stats.items()}
        final_ps = {phrase_lemma_str: {'c': d['count'], 'f': dict(d['forms'])} for phrase_lemma_str, d in
                    self.phrase_stats.items()}
        total_matches = sum(d['count'] for d in self.word_stats.values()) + sum(
            d['count'] for d in self.phrase_stats.values())

        return {'word_stats': final_ws, 'phrase_stats': final_ps, 'total_matches': total_matches}
        #[end]

    def save(self, output_path: str) -> None:
        self.document.save(output_path)
