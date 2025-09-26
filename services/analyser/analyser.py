import docx
from typing import List

from docx.enum.text import WD_COLOR_INDEX
from docx.text.paragraph import Paragraph
from docx.text.hyperlink import Hyperlink
from docx.table import Table
from docx.oxml.ns import qn
from docx.oxml import OxmlElement, CT_R, CT_Text, CT_P, CT_Hyperlink  # CT_R is Complex Type _ Run

from services.analyser import AnalyseData
from services.common_docx import tokenize_paragraph_universal, tokenize_text
from services.highlight_service import _find_matches_in_paragraph_tokens
from utils.timeit import timeit
from copy import deepcopy


class Analyser:
    document: docx.Document
    analyse_data: AnalyseData

    def __init__(self, document: docx.Document | str):
        if isinstance(document, str):
            document = docx.Document(document)

        self.document = document

    def set_analyse_data(self, analyse_data):
        self.analyse_data = analyse_data

    @staticmethod
    def __clone_run(source_run: CT_R, new_text: str, highlight=False) -> CT_R:
        new_run: CT_R = OxmlElement('w:r')

        for child in source_run:
            new_run.append(deepcopy(child))

        # update text
        text_element: CT_Text = new_run.find(qn('w:t'))
        text_element.text = new_text
        text_element.set(qn('xml:space'), 'preserve')

        # update highlight
        if highlight:
            rPr = new_run.find(qn('w:rPr'))

            if rPr is None:
                rPr = OxmlElement('w:rPr')
                new_run.insert(0, rPr)

            highlight = OxmlElement('w:highlight')
            highlight.set(qn('w:val'), 'green')  # зеленая подсветка (BRIGHT_GREEN)
            rPr.append(highlight)

        return new_run

    # find and split run on three new: before, middle and after parts where middle arranged with supplied indices
    @staticmethod
    def __isolate_new_run_xml(batch, phrase_start_index, phrase_end_index) -> List[CT_R]:
        # after Run split operation, batch need to be updated
        new_batch = []

        # [start] find run by char index
        proceed_chars = 0

        for run_source in batch:
            source_run_text_element: CT_Text = run_source.find(qn('w:t'))

            # run has no text
            if source_run_text_element == None:
                new_batch.append(run_source)
                continue

            text: str = source_run_text_element.text
            run_start_index = proceed_chars
            run_end_index = run_start_index + len(text)

            if run_start_index <= phrase_start_index <= run_end_index:
                match_length: int = phrase_end_index - phrase_start_index
                run_relative_char_start_index: int = phrase_start_index - proceed_chars
                run_relative_char_end_index: int = run_relative_char_start_index + match_length

                # [start] split run text on three new parts: before, match and after
                part_before_match: str = text[:run_relative_char_start_index]
                part_match: str = text[run_relative_char_start_index:run_relative_char_end_index]
                part_after_match: str = text[run_relative_char_end_index:]

                # needs to avoid cases "This is an apple" -> "This is anapple"
                source_run_text_element.set(qn('xml:space'), 'preserve')

                run_match = Analyser.__clone_run(run_source, part_match, True)
                run_after = Analyser.__clone_run(run_source, part_after_match)
                source_run_text_element.text = part_before_match
                run_source.addnext(run_match)
                run_match.addnext(run_after)
                # [end]

                new_batch.append(run_source)
                new_batch.append(run_match)
                new_batch.append(run_after)
            else:
                new_batch.append(run_source)

            proceed_chars += len(text)
        # [end]

        return new_batch

    def __process_batch(self, batch: List[CT_R]) -> None:
        text = ''

        # [start] find matches in concatenated batch text
        for run in batch:
            text += run.text

        tokens = tokenize_text(text)
        matches = _find_matches_in_paragraph_tokens(
            tokens,
            self.analyse_data.lemmas,
            self.analyse_data.stems,
            {}
        )
        # [end]

        for match in matches:
            start_token_idx = match['start_token_idx']
            end_token_idx = match['end_token_idx']

            for i in range(start_token_idx, end_token_idx + 1):
                token = tokens[i]
                batch = self.__isolate_new_run_xml(batch, token['start'], token['end'])

    @staticmethod
    def __split_on_batches(element: CT_P | CT_Hyperlink) -> List[List[CT_R]]:
        batches = []
        batch_runs = []
        qn_ct_r = qn('w:r')

        # [start] split paragraph text on batches of runs
        for child in element:
            if child.tag == qn_ct_r:
                batch_runs.append(child)
            elif len(batch_runs) > 0:
                batches.append(batch_runs)
                batch_runs = []

        if len(batch_runs) > 0:
            batches.append(batch_runs)
        # [end]

        return batches

    def __analyse_paragraph(self, paragraph: Paragraph):
        batches = Analyser.__split_on_batches(paragraph._element)

        for batch in batches:
            self.__process_batch(batch)

        hyperlinks: List[Hyperlink] = paragraph.hyperlinks

        for link in hyperlinks:
            self.__analyze_link(link._element)

    def __analyze_link(self, link: CT_Hyperlink):
        batches = Analyser.__split_on_batches(link)

        for batch in batches:
            self.__process_batch(batch)

        return None

    @timeit
    def analyse_and_highlight(self):
        paragraphs: List[Paragraph] = self.document.paragraphs

        for paragraph in paragraphs:
            self.__analyse_paragraph(paragraph)

        tables: List[Table] = self.document.tables

        for table in tables:
            for row in table.rows:
                for cell in row.cells:
                    paragraphs: List[Paragraph] = cell.paragraphs

                    for paragraph in paragraphs:
                        self.__analyse_paragraph(paragraph)

    def save(self, output_path: str) -> None:
        self.document.save(output_path)
