# -*- coding: utf-8 -*-
# --- START OF FILE highlight_service.py ---
import json
import os
import time
from collections import Counter, defaultdict
import copy
import re
import itertools

from flask import current_app

from services.common_docx import _safe_copy_paragraph_format, _apply_run_style, tokenize_paragraph_universal
import docx
from docx.enum.text import WD_COLOR_INDEX
import pandas as pd
import fitz

from services.document_service import extract_all_logical_words_from_pdf, is_predominantly_non_alphabetic
from services.docx_cache import DocxCache

# --- АДАПТИРОВАННЫЕ Импорты из pymorphy_service ---
from services.pymorphy_service import (
    _get_lemma, _get_stem,  # Базовые функции остаются
    reset_caches,  # <--- НОВОЕ ИМЯ для сброса
    WORD_TOKENIZE_PATTERN,  # Используется в _find_matches_in_paragraph_tokens через _get_lemma
    CYRILLIC_PATTERN
)

# OCR Service imports
from services.ocr_service import (
    ocr_page,
    OCR_LANGUAGES,
    OCR_DPI
)
from utils.timeit import timeit

# --- Константы ---
# (Без изменений)
WORDS_EXTRACT_PATTERN = re.compile(r'[a-zA-Zа-яА-ЯёЁ]+', re.UNICODE)
HIGHLIGHT_COLOR_DOCX = WD_COLOR_INDEX.BRIGHT_GREEN
HIGHLIGHT_COLOR_PDF = (0.0, 1.0, 0.0)
# PUNCT_STRIP_PATTERN = re.compile(r"^[^\w\s]+|[^\w\s]+$", re.UNICODE)
PUNCT_STRIP_PATTERN = re.compile(r"^[^\w\s]+|[^\w\s]+$", re.UNICODE)
MIN_OCR_CONFIDENCE_HIGHLIGHT = 40
HYPHEN_CHARS = ('-', '\u00AD')
MAX_LINE_JUMP_MERGE = 150
HORIZONTAL_INDENT_THRESHOLD = 200
MIN_CONF_FOR_MERGE = 5
USE_STEM_FALLBACK = True
STOP_WORDS_RU = {
    # Союзы (Conjunctions)
    "и", "а", "но", "да", "или", "либо", "то", "не то", "тоже", "также",
    "зато", "однако", "же", "что", "чтобы", "как", "будто", "словно",
    "если", "когда", "пока", "едва", "лишь", "потому что", "так как",
    "ибо", "оттого что", "поскольку", "хотя", "хоть", "несмотря на то что",
    "пускай", "пусть", "словно", "точно", "чем", "так что", "поэтому", "причем", "притом",

    # Предлоги (Prepositions)
    "в", "на", "с", "о", "у", "к", "по", "за", "из", "от", "до", "под",
    "над", "при", "без", "для", "про", "об", "обо", "со", "ко", "из-за",
    "из-под", "через", "перед", "между", "среди", "возле", "около",
    "вокруг", "вдоль", "вместо", "внутри", "вне", "кроме", "помимо",
    "сверх", "сквозь", "согласно", "благодаря", "вопреки", "навстречу",
    "ввиду", "вследствие", "наподобие", "насчет", "спустя",

    # Из вашего примера (частицы и др.)
    "не", "бы", "ли", "её"

    # Буквы русского алфавита
                      "а", "б", "в", "г", "д", "е", "ё", "ж", "з", "и", "й", "к", "л", "м",
    "н", "о", "п", "р", "с", "т", "у", "ф", "х", "ц", "ч", "ш", "щ", "ъ",
    "ы", "ь", "э", "ю", "я"
}

STOP_WORDS_EN = {
    # Conjunctions
    "and", "but", "or", "nor", "for", "so", "yet",  # Coordinating
    "after", "although", "as", "because", "before", "if", "once",
    "since", "than", "that", "though", "till", "unless", "until",
    "when", "whenever", "where", "whereas", "wherever", "whether", "while",

    # Prepositions
    "about", "above", "across", "after", "against", "along", "amid", "among",
    "around", "at", "atop", "before", "behind", "below", "beneath", "beside",
    "between", "beyond", "by", "concerning", "despite", "down", "during",
    "except", "for", "from", "in", "inside", "into", "like", "near", "of",
    "off", "on", "onto", "out", "outside", "over", "past", "regarding",
    "since", "through", "throughout", "to", "toward", "under", "underneath",
    "until", "unto", "up", "upon", "with", "within", "without",

    # From your example (articles, verbs)
    "a", "an", "the", "is", "am", "are",

    # Letters of the English alphabet (lowercase)
    "a", "b", "c", "d", "e", "f", "g", "h", "i", "j", "k", "l", "m",
    "n", "o", "p", "q", "r", "s", "t", "u", "v", "w", "x", "y", "z"

}


# --- Функция многоступенчатого поиска слова ---
# (Без изменений, использует _get_lemma, _get_stem)
def find_word_match_multi_stage(
        doc_word_text,
        search_lemmas_set,
        search_stems_set
):
    if not doc_word_text or not isinstance(doc_word_text, str):
        return False, None, None
    doc_word_lower = doc_word_text.lower().strip()

    # --- НОВАЯ ЛОГИКА ОЧИСТКИ ---
    match = PUNCT_STRIP_PATTERN.sub('', doc_word_lower)
    # if match:
    #     # Берем захваченную группу (текст между не-буквенно-цифровыми символами)
    #     word_for_morph = match.group(1)
    # else:
    #     # Если паттерн не сработал (например, строка состоит только из пунктуации),
    #     # используем результат strip(). Если и он пуст, это будет обработано ниже.
    #     word_for_morph = doc_word_lower

    # Дополнительно убедимся, что остались только буквы (на случай если PUNCT_STRIP_PATTERN оставил что-то не то)
    # Эта строка может быть избыточной, если PUNCT_STRIP_PATTERN работает идеально, но добавляет надежности.
    # print(f"DEBUG find_word_match_multi_stage: After PUNCT_STRIP_PATTERN.sub = '{match}'")
    final_letters_match = WORDS_EXTRACT_PATTERN.search(match)
    if final_letters_match:
        word_for_morph = final_letters_match.group(0)
    else:
        # Если после всех очисток не осталось букв (например, было "*,")
        word_for_morph = ""  # или None, чтобы не передавать пустую строку в лемматизатор

    # print(f"DEBUG find_word_match_multi_stage: word_for_morph = '{word_for_morph}'")
    # --- КОНЕЦ НОВОЙ ЛОГИКИ ОЧИСТКИ ---

    if not word_for_morph:  # Проверяем, что после очистки что-то осталось
        # print(f"DEBUG find_word_match_multi_stage: word_for_morph is empty. No match.")
        return False, None, None

    # Используем новые обертки, которые сами кэшируют
    doc_lemma = _get_lemma(word_for_morph)  # use_cache больше не нужен
    doc_stem = None
    if USE_STEM_FALLBACK:
        doc_stem = _get_stem(word_for_morph)  # use_cache больше не нужен
    if not doc_lemma:
        doc_lemma = doc_word_lower
    if doc_lemma in search_lemmas_set:
        is_russian = bool(CYRILLIC_PATTERN.search(doc_lemma or ''))  # Определяем язык по лемме
        stop_words_set = STOP_WORDS_RU if is_russian else STOP_WORDS_EN
        if doc_lemma in stop_words_set:
            return False, None, None
        return True, doc_lemma, 'lemma'
    if USE_STEM_FALLBACK and doc_stem and doc_stem in search_stems_set:
        is_russian = bool(CYRILLIC_PATTERN.search(doc_lemma or ''))
        stop_words_set = STOP_WORDS_RU if is_russian else STOP_WORDS_EN
        if doc_lemma in stop_words_set:
            return False, None, None  # Игнорируем
        return True, doc_lemma, 'stem'
    return False, None, None


# --- Функция воссоздания параграфа DOCX с подсветкой ---
# (Без изменений)
def _reconstruct_paragraph_with_highlighting(
        new_paragraph,
        old_paragraph,
        tokens,
        matches,
        word_stats,
        phrase_stats,
        document_cache: DocxCache
):

    try:
        style_name = old_paragraph.style.name
        new_paragraph.style = document_cache.getStyle(style_name)
    except Exception:
        pass

    try:
        _safe_copy_paragraph_format(
            old_paragraph.paragraph_format,
            new_paragraph.paragraph_format
        )
    except Exception:
        pass

    if not tokens: return

    current_token_idx = 0
    source_runs = list(old_paragraph.runs)
    run_char_positions = []
    current_char_pos = 0

    for run in source_runs:
        run_len = len(run.text)
        run_char_positions.append((current_char_pos, current_char_pos + run_len, run))
        current_char_pos += run_len

    def find_source_run(char_idx_in_para):
        for start, end, run_obj in run_char_positions:
            if start <= char_idx_in_para < end: return run_obj
        return source_runs[-1] if source_runs else None

    # fast
    for match in matches:
        match_start_idx = match['start_token_idx']
        match_end_idx = match['end_token_idx']
        match_type_val = match['type']

        if match_start_idx > current_token_idx:
            for i in range(current_token_idx, match_start_idx):
                token = tokens[i];
                run_style = find_source_run(token['start'])
                nr = new_paragraph.add_run(token['text'])
                if run_style: _apply_run_style(run_style, nr, document_cache, copy_highlight=True)

        text_parts = []

        for i in range(match_start_idx, match_end_idx + 1):
            token = tokens[i];
            text_parts.append(token['text'])
            run_style = find_source_run(token['start'])
            nr = new_paragraph.add_run(token['text'])

            if run_style: _apply_run_style(run_style, nr, document_cache, copy_highlight=True)

            try:
                nr.font.highlight_color = HIGHLIGHT_COLOR_DOCX
            except Exception:
                pass

        found_text = "".join(text_parts).strip()

        if match_type_val == 'phrase':
            lemma_key = match['lemma_key'];
            stats = phrase_stats[" ".join(lemma_key)]
            stats['count'] += 1;
            stats['forms'][found_text] += 1
        elif match_type_val == 'word':
            lemma_key = match['lemma_key']

            if lemma_key:
                stats = word_stats[lemma_key]
                stats['count'] += 1
                stats['forms'][found_text.lower()] += 1

        current_token_idx = match_end_idx + 1

    # add not highlight runs
    for i in range(current_token_idx, len(tokens)):
        token = tokens[i]
        source_run = find_source_run(token['start'])

        target_run = new_paragraph.add_run(token['text'])

        if source_run:
            _apply_run_style(
                source_run,
                target_run,
                document_cache,
                copy_highlight=True
            )


# --- Функция поиска совпадений в токенах параграфа DOCX ---
# (Изменено использование _get_lemma)
def _find_matches_in_paragraph_tokens(
        tokens,
        search_lemmas_set,
        search_stems_set,
        search_phrase_lemmas_map  # Эта карта теперь приходит готовой
):
    matches = []
    matched_token_indices = set()
    num_tokens = len(tokens)

    # 1. Поиск ФРАЗ
    if search_phrase_lemmas_map:
        # Карта уже содержит tuple лемм как ключ
        # Сортируем ключи (кортежи лемм) по длине для приоритета длинных фраз
        sorted_phrase_keys = sorted(search_phrase_lemmas_map.keys(), key=len, reverse=True)

        for phrase_lemma_tuple in sorted_phrase_keys:
            phrase_len = len(phrase_lemma_tuple)
            if phrase_len < 2: continue

            token_idx = 0
            while token_idx <= num_tokens - phrase_len:
                if token_idx in matched_token_indices:
                    token_idx += 1;
                    continue

                # Собираем окно токенов-слов
                window_word_tokens = []
                current_sub_idx = token_idx
                needed_words = phrase_len
                start_actual_token_idx = -1
                end_actual_token_idx = -1

                while needed_words > 0 and current_sub_idx < num_tokens:
                    token = tokens[current_sub_idx]
                    if start_actual_token_idx == -1 and token['type'] == 'word':
                        start_actual_token_idx = current_sub_idx
                    if token['type'] == 'word':
                        window_word_tokens.append(token)
                        end_actual_token_idx = current_sub_idx
                        needed_words -= 1
                    elif start_actual_token_idx != -1:
                        end_actual_token_idx = current_sub_idx

                    if current_sub_idx in matched_token_indices:
                        token_idx = start_actual_token_idx + 1 if start_actual_token_idx != -1 else current_sub_idx + 1
                        window_word_tokens = [];
                        break
                    current_sub_idx += 1

                if not window_word_tokens or len(window_word_tokens) != phrase_len:
                    if start_actual_token_idx != -1 and token_idx <= start_actual_token_idx:
                        token_idx = start_actual_token_idx + 1
                    else:
                        token_idx += 1
                    continue

                # Проверяем пересечение со всем диапазоном токенов
                current_range_indices = set(range(start_actual_token_idx, end_actual_token_idx + 1))
                if not current_range_indices.intersection(matched_token_indices):
                    # Лемматизируем слова из окна
                    window_lemmas = []
                    valid_window = True
                    for tok in window_word_tokens:
                        lemma = _get_lemma(tok['text'])  # Используем новую обертку
                        if lemma is None:  # Если слово не удалось обработать
                            valid_window = False;
                            break
                        window_lemmas.append(lemma)

                    if valid_window and tuple(window_lemmas) == phrase_lemma_tuple:
                        # Найдена фраза!
                        matches.append({
                            'type': 'phrase',
                            'start_token_idx': start_actual_token_idx,
                            'end_token_idx': end_actual_token_idx,
                            'lemma_key': phrase_lemma_tuple
                        })
                        matched_token_indices.update(current_range_indices)
                        token_idx = end_actual_token_idx + 1;
                        continue  # Перескакиваем

                # Сдвигаем начало поиска
                token_idx = start_actual_token_idx + 1 if start_actual_token_idx != -1 else token_idx + 1

    # 2. Поиск ОТДЕЛЬНЫХ СЛОВ
    for i, token in enumerate(tokens):
        if i in matched_token_indices: continue
        if token['type'] == 'word':
            # Используем find_word_match_multi_stage, который внутри вызывает _get_lemma/_get_stem
            is_match, match_lemma, match_type_detail = find_word_match_multi_stage(
                token['text'],
                search_lemmas_set,
                search_stems_set
            )
            if is_match:
                matches.append({
                    'type': 'word',
                    'start_token_idx': i,
                    'end_token_idx': i,
                    'lemma_key': match_lemma,
                    'word_match_type': match_type_detail  # 'lemma' или 'stem'
                })
                matched_token_indices.add(i)
    matches.sort(key=lambda m: m['start_token_idx'])
    return matches


# --- Основная функция analyze_and_highlight_docx ---
@timeit
def analyze_and_highlight_docx(
        source_path,
        search_data,
        search_phrase_lemmas_map,
        output_path
):
    search_lemmas_set = search_data.get('lemmas', set())
    search_stems_set = search_data.get('stems', set())
    word_stats = defaultdict(lambda: {'count': 0, 'forms': Counter()})
    phrase_stats = defaultdict(lambda: {'count': 0, 'forms': Counter()})
    total_matches_count = 0

    try:
        try:
            source_doc = docx.Document(source_path)
        except Exception:
            return None

        result_doc = docx.Document()

        # ... (Копирование стилей и секций без изменений) ...
        for style in source_doc.styles:
            if style.type == docx.enum.style.WD_STYLE_TYPE.PARAGRAPH or \
                    style.type == docx.enum.style.WD_STYLE_TYPE.CHARACTER:
                try:
                    target_style = result_doc.styles.add_style(style.name, style.type)

                    if style.base_style and style.base_style.name in result_doc.styles:
                        target_style.base_style = result_doc.styles[style.base_style.name]

                    if hasattr(style, 'font') and style.font:
                        if hasattr(target_style, 'font') and target_style.font:
                            target_style.font.name = style.font.name
                            target_style.font.size = style.font.size

                    if hasattr(style, 'paragraph_format') and style.paragraph_format:
                        if hasattr(target_style, 'paragraph_format') and target_style.paragraph_format:
                            target_style.paragraph_format.alignment = style.paragraph_format.alignment

                except Exception:
                    pass

        for section_idx, section in enumerate(source_doc.sections):
            target_section = result_doc.sections[section_idx if section_idx < len(result_doc.sections) else -1]
            target_section.start_type = section.start_type;
            target_section.orientation = section.orientation
            target_section.page_width = section.page_width;
            target_section.page_height = section.page_height
            target_section.left_margin = section.left_margin;
            target_section.right_margin = section.right_margin
            target_section.top_margin = section.top_margin;
            target_section.bottom_margin = section.bottom_margin
            target_section.header_distance = section.header_distance;
            target_section.footer_distance = section.footer_distance

            if section_idx < len(source_doc.sections) - 1 and len(result_doc.sections) == section_idx + 1:
                result_doc.add_section(source_doc.sections[section_idx + 1].start_type)

        body_elements = []

        for element in source_doc.element.body:
            if isinstance(element, docx.oxml.text.paragraph.CT_P):
                body_elements.append({'type': 'p', 'el': element})
            elif isinstance(element, docx.oxml.table.CT_Tbl):
                body_elements.append({'type': 'tbl', 'el': element})

        # save getter results (slow operations)
        paragraphs = source_doc.paragraphs
        tables = source_doc.tables
        document_cache: DocxCache = DocxCache(source_doc)

        # Обработка элементов
        for item in body_elements:
            el_type, element = item['type'], item['el']
            if el_type == 'p':
                source_p = next((p for p in paragraphs if p._element is element), None)

                if source_p is None: continue
                new_p = result_doc.add_paragraph()
                try:
                    tokens = tokenize_paragraph_universal(source_p)

                    if not tokens:  # Пустой параграф
                        if source_p.style and hasattr(source_p.style, 'name'):
                            try:
                                style_name = source_p.style.name
                                if style_name in result_doc.styles: new_p.style = result_doc.styles[style_name]
                            except Exception:
                                pass

                        _safe_copy_paragraph_format(source_p.paragraph_format, new_p.paragraph_format)
                    else:  # Непустой параграф
                        # Используем переданные search_data и search_phrase_lemmas_map
                        matches = _find_matches_in_paragraph_tokens(
                            tokens,
                            search_lemmas_set,
                            search_stems_set,
                            search_phrase_lemmas_map
                        )
                        _reconstruct_paragraph_with_highlighting(
                            new_p,
                            source_p,
                            tokens,
                            matches,
                            word_stats,
                            phrase_stats,
                            document_cache
                        )
                        total_matches_count += len(matches)
                except Exception as e_para:
                    try:  # Fallback
                        new_p.text = source_p.text
                        if source_p.style and hasattr(source_p.style, 'name'):
                            try:
                                style_name = source_p.style.name
                                if style_name in result_doc.styles: new_p.style = result_doc.styles[style_name]
                            except Exception:
                                pass
                        if hasattr(source_p, 'paragraph_format'):
                            _safe_copy_paragraph_format(source_p.paragraph_format, new_p.paragraph_format)
                    except Exception:
                        result_doc.add_paragraph(f"[Ошибка обработки параграфа: {e_para}]")
            elif el_type == 'tbl':
                source_t = next((t for t in tables if t._element is element), None)
                if source_t is None: continue
                try:
                    new_t = result_doc.add_table(rows=len(source_t.rows), cols=len(source_t.columns))
                    # ... (копирование стилей и свойств таблицы) ...
                    if source_t.style and hasattr(source_t.style, 'name'):
                        try:
                            style_name = source_t.style.name
                            if style_name in result_doc.styles: new_t.style = result_doc.styles[style_name]
                        except Exception:
                            pass
                    try:
                        new_t.autofit = source_t.autofit
                    except AttributeError:
                        pass
                    try:
                        new_t.alignment = source_t.alignment
                    except AttributeError:
                        pass

                    for i, row in enumerate(source_t.rows):
                        if i >= len(new_t.rows): continue
                        new_row_cells = new_t.rows[i].cells
                        for j, cell in enumerate(row.cells):
                            if j >= len(new_row_cells): continue
                            target_c = new_row_cells[j];
                            source_c = cell
                            for p_idx in range(len(target_c.paragraphs) - 1, -1, -1):
                                p_to_remove = target_c.paragraphs[p_idx]
                                p_to_remove._element.getparent().remove(p_to_remove._element)
                            try:
                                target_c.vertical_alignment = source_c.vertical_alignment
                            except Exception:
                                pass

                            for cell_p in source_c.paragraphs:
                                new_cp = target_c.add_paragraph()
                                try:
                                    tokens = tokenize_paragraph_universal(cell_p)
                                    if not tokens:  # Пустой параграф ячейки
                                        if cell_p.style and hasattr(cell_p.style, 'name'):
                                            try:
                                                style_name = cell_p.style.name
                                                if style_name in result_doc.styles: new_cp.style = result_doc.styles[
                                                    style_name]
                                            except Exception:
                                                pass
                                        if hasattr(cell_p, 'paragraph_format'):
                                            _safe_copy_paragraph_format(cell_p.paragraph_format,
                                                                        new_cp.paragraph_format)
                                    else:  # Непустой параграф ячейки
                                        # Используем переданные данные
                                        matches = _find_matches_in_paragraph_tokens(
                                            tokens, search_lemmas_set, search_stems_set, search_phrase_lemmas_map
                                        )
                                        _reconstruct_paragraph_with_highlighting(
                                            new_cp,
                                            cell_p,
                                            tokens,
                                            matches,
                                            word_stats,
                                            phrase_stats,
                                            document_cache
                                        )
                                        total_matches_count += len(matches)
                                except Exception as e_cp:
                                    try:  # Fallback ячейки
                                        new_cp.text = cell_p.text
                                        if cell_p.style and hasattr(cell_p.style, 'name'):
                                            try:
                                                style_name = cell_p.style.name
                                                if style_name in result_doc.styles: new_cp.style = result_doc.styles[
                                                    style_name]
                                            except Exception:
                                                pass
                                        if hasattr(cell_p, 'paragraph_format'):
                                            _safe_copy_paragraph_format(cell_p.paragraph_format,
                                                                        new_cp.paragraph_format)
                                    except Exception:
                                        target_c.add_paragraph(f"[Ошибка параграфа ячейки {i},{j}: {e_cp}]")
                except Exception as e_tbl:
                    result_doc.add_paragraph(f"[Ошибка копирования таблицы: {e_tbl}]")

        result_doc.save(output_path)
        final_ws = {l: {'c': d['count'], 'f': dict(d['forms'])} for l, d in word_stats.items()}
        final_ps = {phrase_lemma_str: {'c': d['count'], 'f': dict(d['forms'])} for phrase_lemma_str, d in
                    phrase_stats.items()}
        # Важно: total_matches_count - это количество совпадений (matches),
        # а не уникальных слов/фраз. Можно вернуть оба или выбрать одно.
        # Вернем total_matches, как и раньше.
        return {'word_stats': final_ws, 'phrase_stats': final_ps, 'total_matches': total_matches_count}
    except FileNotFoundError:
        return None
    except Exception:
        if os.path.exists(output_path):
            try:
                os.remove(output_path)
            except OSError:
                pass
        return None


# --- Унифицированная функция analyze_and_highlight_pdf ---
def analyze_and_highlight_pdf(
        source_path,
        search_data,  # <--- Ожидаем результат get_highlight_search_data()
        search_phrase_lemmas_map,  # <--- Ожидаем результат get_highlight_phrase_map()
        output_path,
        use_ocr=False
):
    # Подготовка данных теперь ВНЕ этой функции
    # reset_caches() # <--- Вызываем новый сброс кэша

    search_lemmas_set = search_data.get('lemmas', set())
    search_stems_set = search_data.get('stems', set())
    # Векторы удалены

    word_stats = defaultdict(lambda: {'count': 0, 'forms': Counter()})
    phrase_stats = defaultdict(lambda: {'count': 0, 'forms': Counter()})
    total_highlight_actions = 0
    doc = None
    try:
        doc = fitz.open(source_path)
    except Exception:
        return None

    all_pages_logical_words_fitz = extract_all_logical_words_from_pdf(source_path)
    if all_pages_logical_words_fitz is None:
        if doc: doc.close(); return None
    if len(doc) != len(all_pages_logical_words_fitz):
        if doc: doc.close(); return None

    # Цикл по страницам
    for page_num, logical_words_on_page_fitz in enumerate(all_pages_logical_words_fitz):
        page = doc.load_page(page_num)
        page_rect_fitz = page.rect
        highlighted_rects_on_page = []
        page_word_candidates = []
        page_word_candidates_ocr_processed = []

        # Сбор кандидатов из Fitz (без изменений)
        for word_data in logical_words_on_page_fitz:
            word_text_fitz = word_data.get('text', '').strip()
            word_rects_fitz = word_data.get('rects', [])
            if not word_text_fitz or not word_rects_fitz: continue
            page_word_candidates.append(
                {'text': word_text_fitz, 'rects': word_rects_fitz, 'source': 'fitz', 'confidence': 100.0})

        # Сбор кандидатов из OCR (без изменений)
        if use_ocr:
            try:
                print(f"Page {page_num + 1}: Calling ocr_page...")
                ocr_data, ocr_matrix, _ = ocr_page(page, languages=OCR_LANGUAGES, dpi=OCR_DPI)
                print(
                    f"Page {page_num + 1}: ocr_page returned. Data is None: {ocr_data is None}, Matrix is None: {ocr_matrix is None}")
                if ocr_data is not None:
                    print(f"Page {page_num + 1}: ocr_data length: {len(ocr_data)}")
                if ocr_data is not None and not ocr_data.empty and ocr_matrix is not None:
                    if ocr_data is not None and not ocr_data.empty:
                        print(f"--- DEBUG: Checking ocr_data before merge (Page {page_num + 1}) ---")
                        # Попробуем найти нужные строки
                        shulman_rows = ocr_data[ocr_data['text'].str.contains("ШУЛЬМАН", na=False)]
                        ekaterina_rows = ocr_data[ocr_data['text'].str.contains("ЕКАТЕРИНА", na=False)]
                        print("Rows containing ШУЛЬМАН:")
                        print(shulman_rows[['level', 'text', 'conf', 'left', 'top', 'width', 'height']])
                        print("Rows containing ЕКАТЕРИНА:")
                        print(ekaterina_rows[['level', 'text', 'conf', 'left', 'top', 'width', 'height']])
                        print("--- END DEBUG ---")
                    processed_ocr_words = _merge_hyphenated_ocr_words(ocr_data)
                    print(
                        f"Page {page_num + 1}: _merge_hyphenated_ocr_words returned {len(processed_ocr_words)} words.")

                    try:
                        inverse_ocr_matrix = ~ocr_matrix;
                        matrix_valid = True
                    except ValueError:
                        matrix_valid = False
                    if matrix_valid:
                        for word_info in processed_ocr_words:
                            ocr_text = word_info['text'];
                            ocr_conf = word_info['conf']
                            original_indices = word_info['original_indices']
                            if ocr_conf < MIN_OCR_CONFIDENCE_HIGHLIGHT: continue
                            ocr_rect_tuples = [];
                            valid_coords_for_word = True
                            for index in original_indices:
                                try:
                                    row = ocr_data.loc[index]
                                    left, top, width, height = int(row['left']), int(row['top']), int(
                                        row['width']), int(row['height'])
                                    if width <= 0 or height <= 0: valid_coords_for_word = False; break
                                    img_rect = fitz.Rect(left, top, left + width, top + height)
                                    pdf_rect = img_rect * inverse_ocr_matrix;
                                    pdf_rect.normalize()
                                    pdf_rect = pdf_rect & page_rect_fitz
                                    if pdf_rect and not pdf_rect.is_empty and pdf_rect.width > 1e-3 and pdf_rect.height > 1e-3:
                                        ocr_rect_tuples.append(pdf_rect.irect)
                                    else:
                                        valid_coords_for_word = False;
                                        break
                                except (KeyError, ValueError, TypeError):
                                    valid_coords_for_word = False;
                                    break
                            if valid_coords_for_word and ocr_rect_tuples:
                                page_word_candidates_ocr_processed.append({
                                    'text': ocr_text, 'rects': ocr_rect_tuples,
                                    'source': 'ocr', 'confidence': ocr_conf
                                })
            except Exception:
                pass

        # Дедупликация кандидатов (без изменений)
        all_page_word_candidates = page_word_candidates + page_word_candidates_ocr_processed
        unique_page_word_candidates = _deduplicate_candidates(all_page_word_candidates)

        # Унифицированный поиск и подсветка СЛОВ (используем переданные данные)
        for candidate in unique_page_word_candidates:
            candidate_text = candidate['text']
            candidate_rects_tuples = candidate['rects']

            # Используем переданные search_lemmas_set и search_stems_set
            is_match, match_lemma, _ = find_word_match_multi_stage(
                candidate_text, search_lemmas_set, search_stems_set
            )
            if is_match:
                candidate_bbox = _get_bounding_rect(candidate_rects_tuples)
                if candidate_bbox is None or candidate_bbox.is_empty: continue
                # ... (проверка пересечения и подсветка без изменений) ...
                overlaps_existing = False;
                check_bbox = candidate_bbox + (-1, -1, 1, 1)
                overlap_threshold_area_ratio = 0.2
                for existing_hl_rect in highlighted_rects_on_page:
                    if check_bbox.intersects(existing_hl_rect):
                        intersection = check_bbox & existing_hl_rect;
                        intersection_area = intersection.get_area()
                        bbox_area = candidate_bbox.get_area()
                        if bbox_area > 1e-5 and (intersection_area / bbox_area > overlap_threshold_area_ratio):
                            overlaps_existing = True;
                            break
                        elif bbox_area <= 1e-5 and intersection_area > 1e-5:
                            overlaps_existing = True;
                            break
                if not overlaps_existing:
                    current_word_rects_highlighted = [];
                    quads_added_count_word = 0
                    for rect_tuple in candidate_rects_tuples:
                        try:
                            rect = fitz.Rect(rect_tuple)
                            if rect.is_empty: continue
                            highlight = page.add_highlight_annot(rect)
                            if highlight:
                                highlight.set_colors(stroke=HIGHLIGHT_COLOR_PDF);
                                highlight.update(opacity=0.4)
                                quads_added_count_word += 1;
                                total_highlight_actions += 1
                                current_word_rects_highlighted.append(rect)
                        except Exception:
                            pass
                    if quads_added_count_word > 0:
                        stats_word = word_stats[match_lemma]
                        stats_word['count'] += 1
                        stats_word['forms'][candidate_text.lower()] += 1
                        highlighted_rects_on_page.extend(current_word_rects_highlighted)

        # Поиск и подсветка ФРАЗ (используем переданную карту)
        fitz_tokens_for_phrases = []
        for word_data in logical_words_on_page_fitz:
            text = word_data.get('text', '').strip()
            rects = word_data.get('rects', [])
            if text and rects:
                fitz_tokens_for_phrases.append({'text': text, 'rects': rects})

        # Используем переданную search_phrase_lemmas_map
        if search_phrase_lemmas_map and fitz_tokens_for_phrases:
            sorted_phrase_keys = sorted(search_phrase_lemmas_map.keys(), key=len, reverse=True)
            for phrase_lemma_tuple in sorted_phrase_keys:
                target_len = len(phrase_lemma_tuple)
                if target_len < 2 or target_len > len(fitz_tokens_for_phrases): continue
                for i in range(len(fitz_tokens_for_phrases) - target_len + 1):
                    window_tokens = fitz_tokens_for_phrases[i: i + target_len]
                    window_lemmas = [];
                    valid_phrase_lemmas = True
                    for tok in window_tokens:
                        lemma = _get_lemma(tok['text'])  # Используем новую обертку
                        if lemma is None: valid_phrase_lemmas = False; break
                        window_lemmas.append(lemma)

                    if valid_phrase_lemmas and tuple(window_lemmas) == phrase_lemma_tuple:
                        # ... (проверка пересечения и подсветка без изменений) ...
                        phrase_text = " ".join(tok['text'] for tok in window_tokens)
                        phrase_rects_tuples = list(itertools.chain.from_iterable(tok['rects'] for tok in window_tokens))
                        phrase_bbox = _get_bounding_rect(phrase_rects_tuples)
                        if phrase_bbox is None or phrase_bbox.is_empty: continue
                        overlaps_existing_highlight = False;
                        check_phrase_bbox = phrase_bbox + (-1, -1, 1, 1)
                        overlap_threshold_area_ratio = 0.2
                        phrase_bbox_area = phrase_bbox.get_area()
                        for existing_hl_rect in highlighted_rects_on_page:
                            if check_phrase_bbox.intersects(existing_hl_rect):
                                intersection = check_phrase_bbox & existing_hl_rect;
                                intersection_area = intersection.get_area()
                                if phrase_bbox_area > 1e-5 and (
                                        intersection_area / phrase_bbox_area > overlap_threshold_area_ratio):
                                    overlaps_existing_highlight = True;
                                    break
                                elif phrase_bbox_area <= 1e-5 and intersection_area > 1e-5:
                                    overlaps_existing_highlight = True;
                                    break
                        if not overlaps_existing_highlight:
                            current_phrase_rects_highlighted = [];
                            quads_added_count_phrase = 0
                            for rect_tuple in phrase_rects_tuples:
                                try:
                                    rect = fitz.Rect(rect_tuple)
                                    if rect.is_empty: continue
                                    highlight_annot = page.add_highlight_annot(rect)
                                    if highlight_annot:
                                        highlight_annot.set_colors(stroke=HIGHLIGHT_COLOR_PDF);
                                        highlight_annot.update(opacity=0.4)
                                        quads_added_count_phrase += 1;
                                        total_highlight_actions += 1
                                        current_phrase_rects_highlighted.append(rect)
                                except Exception:
                                    pass
                            if quads_added_count_phrase > 0:
                                phrase_key_str = " ".join(phrase_lemma_tuple)
                                stats_phrase = phrase_stats[phrase_key_str]
                                stats_phrase['count'] += 1;
                                stats_phrase['forms'][phrase_text.strip()] += 1
                                highlighted_rects_on_page.extend(current_phrase_rects_highlighted)
    # Завершение (без изменений)
    total_matches_combined = sum(d['count'] for d in word_stats.values()) + sum(
        d['count'] for d in phrase_stats.values())
    if total_highlight_actions > 0:
        try:
            doc.save(output_path, garbage=4, deflate=True, clean=True)
        except Exception:
            if doc: doc.close(); return None
    else:
        if os.path.exists(output_path):
            try:
                os.remove(output_path)
            except Exception:
                pass
    if doc:
        try:
            doc.close()
        except Exception:
            pass
    final_ws = {l: {'c': d['count'], 'f': dict(d['forms'])} for l, d in word_stats.items()}
    final_ps = {phrase_lemma_str: {'c': d['count'], 'f': dict(d['forms'])} for phrase_lemma_str, d in
                phrase_stats.items()}
    # Возвращаем total_matches_combined для консистентности с DOCX
    return {'word_stats': final_ws, 'phrase_stats': final_ps, 'total_matches': total_matches_combined}


# --- Вспомогательные функции для PDF ---
# (Без изменений: _get_bounding_rect, _rects_overlap_significantly, _deduplicate_candidates, _merge_hyphenated_ocr_words)
def _get_bounding_rect(rects_list_of_tuples):
    if not rects_list_of_tuples: return None
    min_x0, min_y0 = float('inf'), float('inf')
    max_x1, max_y1 = float('-inf'), float('-inf')
    valid_rect_found = False
    for r_tuple in rects_list_of_tuples:
        try:
            x0, y0, x1, y1 = r_tuple
            if x0 < x1 and y0 < y1:
                min_x0 = min(min_x0, x0);
                min_y0 = min(min_y0, y0)
                max_x1 = max(max_x1, x1);
                max_y1 = max(max_y1, y1)
                valid_rect_found = True
        except (TypeError, ValueError):
            continue
    if not valid_rect_found: return None
    return fitz.Rect(min_x0, min_y0, max_x1, max_y1)


def _rects_overlap_significantly(rects1_list_tuples, rects2_list_tuples, threshold=0.5):
    bbox1 = _get_bounding_rect(rects1_list_tuples)
    bbox2 = _get_bounding_rect(rects2_list_tuples)
    if bbox1 is None or bbox2 is None or bbox1.is_empty or bbox2.is_empty: return False
    intersection_rect = bbox1 & bbox2
    if intersection_rect.is_empty: return False
    area1 = bbox1.get_area();
    area2 = bbox2.get_area();
    intersection_area = intersection_rect.get_area()
    if area1 <= 1e-9 or area2 <= 1e-9: return False
    smaller_area = min(area1, area2)
    overlap_ratio = intersection_area / smaller_area
    return overlap_ratio >= threshold


def _deduplicate_candidates(candidates):
    if not candidates: return []
    unique_candidates = [];
    processed_indices = set()
    candidates.sort(key=lambda c: (c['source'] != 'fitz', -float(c.get('confidence', -1.0))))
    for i in range(len(candidates)):
        if i in processed_indices: continue
        current_candidate = candidates[i];

        # --- ЛОГ ДЛЯ ИНТЕРЕСУЮЩЕГО НАС СЛОВА ---
        is_target_word = "ШУЛЬМАН" in current_candidate.get('text', '').upper()
        if is_target_word:
            print(
                f"DEBUG _deduplicate: Processing current_candidate (i={i}): '{current_candidate.get('text')}', Source: {current_candidate.get('source')}, Conf: {current_candidate.get('confidence')}")
        # --- КОНЕЦ ЛОГА ---

        overlapping_group = [current_candidate];
        indices_in_group = {i}
        for j in range(i + 1, len(candidates)):
            if j in processed_indices: continue
            other_candidate = candidates[j]
            if _rects_overlap_significantly(current_candidate['rects'], other_candidate['rects']):
                # --- ЛОГ ПЕРЕСЕЧЕНИЯ ---
                if is_target_word or "ШУЛЬМАН" in other_candidate.get('text',
                                                                      '').upper():  # Логируем, если один из них - наша цель
                    print(
                        f"DEBUG _deduplicate:   Overlap found between i={i} ('{current_candidate.get('text')}') and j={j} ('{other_candidate.get('text')}', Src: {other_candidate.get('source')}, Conf: {other_candidate.get('confidence')})")
                # --- КОНЕЦ ЛОГА ---

                overlapping_group.append(other_candidate);
                indices_in_group.add(j)
        winner = overlapping_group[0]

        # --- ЛОГ РЕШЕНИЯ ПО ГРУППЕ ---
        if is_target_word:  # Если текущий кандидат - наша цель
            if winner == current_candidate:
                print(f"DEBUG _deduplicate:   '{current_candidate.get('text')}' (i={i}) WON in its group.")
            else:
                print(
                    f"DEBUG _deduplicate:   '{current_candidate.get('text')}' (i={i}) LOST to '{winner.get('text')}' (Src: {winner.get('source')}, Conf: {winner.get('confidence')}) in its group.")
            print(f"DEBUG _deduplicate:   Group members for '{current_candidate.get('text')}':")
            for member_idx, member in enumerate(overlapping_group):
                print(
                    f"DEBUG _deduplicate:     Member {member_idx}: '{member.get('text')}', Src: {member.get('source')}, Conf: {member.get('confidence')}")
        # --- КОНЕЦ ЛОГА ---

        unique_candidates.append(winner)
        processed_indices.update(indices_in_group)
    return unique_candidates


def _merge_hyphenated_ocr_words(ocr_df):
    if ocr_df is None or ocr_df.empty: return []
    required_cols = ['level', 'text', 'conf', 'left', 'top', 'width', 'height']
    if not all(col in ocr_df.columns for col in required_cols):
        return [{'text': row.get('text', ''), 'conf': float(row.get('conf', -1.0)),
                 'level': int(row.get('level', -1)), 'is_merged': False,
                 'original_indices': [index]}
                for index, row in ocr_df.iterrows() if int(row.get('level', -1)) == 5]
    words_df = ocr_df[
        (ocr_df['level'].astype(int) == 5) & (pd.to_numeric(ocr_df['conf'], errors='coerce') >= MIN_CONF_FOR_MERGE) & (
            ocr_df['text'].notna()) & (ocr_df['text'].astype(str).str.strip() != '')].copy()
    if words_df.empty: return []
    for col in ['left', 'top', 'width', 'height', 'conf']:
        words_df[col] = pd.to_numeric(words_df[col], errors='coerce')
    words_df.dropna(subset=['left', 'top', 'width', 'height', 'conf'], inplace=True)
    for col in ['left', 'top', 'width', 'height']: words_df[col] = words_df[col].astype(int)
    words_df['conf'] = words_df['conf'].astype(float)
    if words_df.empty: return []
    words_df = words_df.sort_values(by=['top', 'left'])
    output_words = [];
    merged_indices = set();
    word_indices = list(words_df.index)
    for i in range(len(word_indices)):
        idx1 = word_indices[i]
        if idx1 in merged_indices: continue
        row1 = words_df.loc[idx1];
        text1 = str(row1['text']).strip()
        # --- ДОБАВИТЬ ЛОГ ЗДЕСЬ ---
        if "ШУЛЬМАН" in text1.upper() or "ЕКАТЕРИНА" in text1.upper():
            print(f"DEBUG _merge: Processing row index {idx1}, text1='{text1}', conf={row1['conf']:.1f}")
        # --- КОНЕЦ ЛОГА ---
        if not text1: continue
        ends_with_hyphen = text1.endswith(HYPHEN_CHARS);
        found_merge = False
        if ends_with_hyphen:
            for j in range(i + 1, len(word_indices)):
                idx2 = word_indices[j]
                if idx2 in merged_indices: continue
                row2 = words_df.loc[idx2];
                text2 = str(row2['text']).strip()
                if not text2: continue
                is_below = row2['top'] > row1['top']
                vertical_dist_ok = abs(row2['top'] - row1['top']) < MAX_LINE_JUMP_MERGE
                is_left_aligned = row2['left'] < HORIZONTAL_INDENT_THRESHOLD
                if is_below and vertical_dist_ok and is_left_aligned:
                    merged_text = text1.rstrip(''.join(HYPHEN_CHARS)) + text2
                    min_confidence = min(row1['conf'], row2['conf'])
                    output_words.append({'text': merged_text, 'conf': min_confidence, 'level': 5, 'is_merged': True,
                                         'original_indices': [idx1, idx2]})
                    merged_indices.add(idx1);
                    merged_indices.add(idx2)
                    found_merge = True;
                    break
        if not found_merge:
            if "ШУЛЬМАН" in text1.upper() or "ЕКАТЕРИНА" in text1.upper():
                print(f"DEBUG _merge: Adding original word to output: text1='{text1}'")
            # --- КОНЕЦ ЛОГА ---
            output_words.append(
                {'text': text1, 'conf': row1['conf'], 'level': 5, 'is_merged': False, 'original_indices': [idx1]})
    return output_words


# --- END OF FILE highlight_service.py ---
